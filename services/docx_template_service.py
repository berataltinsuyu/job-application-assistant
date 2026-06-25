from copy import deepcopy
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


DOCX_TEMPLATE_DIR = Path("templates/docx")
GENERATED_TEMPLATE_DIR = DOCX_TEMPLATE_DIR / "generated"

DOCX_TEMPLATE_CATALOG = [
    {
        "template_id": "ats_classic_docx",
        "display_name": "ATS Classic DOCX",
        "description": "Compact traditional ATS-friendly DOCX with clear section hierarchy and minimal styling.",
        "best_for": "Conservative applications, banks, enterprise, corporate roles, ATS-heavy roles",
        "visual_style": "Compact, traditional, text-first",
        "layout": "One-column",
        "strengths": "High readability, compact spacing, conventional section hierarchy",
        "cautions": "Less visually expressive than modern template",
        "recommended_for": "Finance, banking, corporate IT, operations, business analysis, backend roles",
        "not_recommended_for": "Design-heavy creative roles",
        "ats_safety_level": "high",
        "visual_density": "medium",
        "supports_docx": True,
        "supports_pdf": False,
        "supports_photo": False,
        "experimental": True,
    },
    {
        "template_id": "ats_modern_docx",
        "display_name": "ATS Modern DOCX",
        "description": "Modern ATS-friendly DOCX with more whitespace, stronger headings, and clean separators.",
        "best_for": "Modern tech roles, startups, software engineering, AI/data roles, product-oriented roles",
        "visual_style": "Modern, clean, more whitespace",
        "layout": "One-column",
        "strengths": "Stronger hierarchy, cleaner spacing, more polished visual impression",
        "cautions": "May use slightly more space than classic",
        "recommended_for": "Software engineering, AI applications, data, product, digital roles",
        "not_recommended_for": "Ultra-conservative one-page corporate submissions where maximum compactness is needed",
        "ats_safety_level": "high",
        "visual_density": "medium",
        "supports_docx": True,
        "supports_pdf": False,
        "supports_photo": False,
        "experimental": True,
    },
    {
        "template_id": "modern_professional",
        "display_name": "Modern Professional",
        "description": "Polished general-purpose ATS-friendly one-column CV with strong name hierarchy, clean section rules, and balanced spacing.",
        "best_for": "Product roles, business analyst roles, corporate IT, business roles, new graduate applications",
        "visual_style": "Professional, centered header, subtle blue-gray rules",
        "layout": "One-column",
        "strengths": "Clear visual hierarchy, polished section separation, compact contact treatment, ATS-readable content order",
        "cautions": "More visual than strict plain ATS; review before highly conservative submissions",
        "recommended_for": "Product, business analyst, corporate IT, operations, junior professional roles",
        "not_recommended_for": "Ultra-compact one-page technical submissions with dense content",
        "ats_safety_level": "high",
        "visual_density": "medium",
        "supports_docx": True,
        "supports_pdf": False,
        "supports_photo": False,
        "experimental": True,
    },
    {
        "template_id": "compact_technical",
        "display_name": "Compact Technical",
        "description": "Compact, technical, ATS-safe CV with tight spacing, strong section rules, and intentional skill grouping.",
        "best_for": "Backend developer, software engineer, data/AI technical roles, ATS-heavy submissions",
        "visual_style": "Compact, technical, rule-led hierarchy",
        "layout": "One-column",
        "strengths": "Dense but readable layout, strong technical keyword visibility, good for project-heavy CVs",
        "cautions": "Less whitespace than modern templates; best for concise bullet content",
        "recommended_for": "Backend, software engineering, API, data, AI, DevOps, technical internship roles",
        "not_recommended_for": "Design-forward or photo-first applications",
        "ats_safety_level": "high",
        "visual_density": "high",
        "supports_docx": True,
        "supports_pdf": False,
        "supports_photo": False,
        "experimental": True,
    },
    {
        "template_id": "visual_photo_optional",
        "display_name": "Visual Photo Optional",
        "description": "ATS-conscious one-column CV with an optional modest header photo for local/Turkish/corporate submissions where photo CVs are acceptable.",
        "best_for": "Turkish/local applications, photo-acceptable corporate submissions, polished visual CVs",
        "visual_style": "Photo-capable header, clean rules, corporate spacing",
        "layout": "One-column with optional header photo",
        "strengths": "Works with or without photo, keeps main content one-column, restrained visual presentation",
        "cautions": "Only include a photo where it is expected or acceptable for the target market",
        "recommended_for": "Turkey/local applications, corporate roles where photo CVs are common",
        "not_recommended_for": "Photo-blind ATS processes, US/UK applications where photos are discouraged",
        "ats_safety_level": "medium",
        "visual_density": "medium",
        "supports_docx": True,
        "supports_pdf": False,
        "supports_photo": True,
        "experimental": True,
    },
]

SECTION_ORDER = [
    "professional_summary",
    "skills",
    "experience",
    "projects",
    "education",
    "certifications",
    "languages",
]

SECTION_TITLES = {
    "professional_summary": "Professional Summary",
    "skills": "Skills",
    "experience": "Experience",
    "projects": "Projects",
    "education": "Education",
    "certifications": "Certifications",
    "languages": "Languages",
}


def ensure_builtin_docx_templates() -> list[dict]:
    """Create local generated sample DOCX files for the built-in renderer catalog."""
    GENERATED_TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    sample_cv = {
        "contact": {
            "full_name": "Sample Candidate",
            "target_title": "Target Role",
            "email": "candidate@example.com",
            "phone": "+1 555 555 5555",
            "location": "Remote",
            "linkedin": "https://linkedin.com/in/sample",
            "github": "https://github.com/sample",
        },
        "professional_summary": "Short ATS-friendly summary generated locally by the built-in DOCX template service.",
        "skills": {"Core": ["Python", "SQL", "REST APIs"], "Tools": ["Git", "Docker"]},
        "experience": [
            {
                "title": "Example Role",
                "company": "Example Organization",
                "start_date": "2025",
                "end_date": "Present",
                "bullets": ["Rendered with python-docx only.", "No external template files are required."],
            }
        ],
        "education": [{"school": "Example University", "degree": "Example Degree"}],
    }
    catalog = get_docx_template_catalog()
    for template in catalog:
        output_path = GENERATED_TEMPLATE_DIR / f"{template['template_id']}_sample.docx"
        if not output_path.exists() or output_path.stat().st_size == 0:
            render_cv_with_docx_template(sample_cv, template["template_id"], str(output_path), {"sample": True})
    return catalog


def get_docx_template_catalog() -> list[dict]:
    return deepcopy(DOCX_TEMPLATE_CATALOG)


def _add_bottom_border(paragraph, color="CCCCCC", size="6", space="4") -> None:
    """Helper to apply a clean bottom border to a paragraph using native word XML."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), str(space))
    bottom.set(qn('w:color'), str(color))
    pBdr.append(bottom)
    pPr.append(pBdr)


def render_cv_with_docx_template(
    structured_cv: dict,
    template_id: str,
    output_path: str,
    metadata: dict | None = None,
    photo_bytes: bytes | None = None,
    photo_filename: str = "",
) -> dict:
    warnings: list[str] = []
    template = _get_template(template_id)
    if not template:
        return {
            "success": False,
            "template_id": template_id,
            "output_path": "",
            "warnings": [f"Unknown DOCX template_id: {template_id}"],
            "message": "Unknown DOCX template.",
        }

    if not isinstance(structured_cv, dict):
        return {
            "success": False,
            "template_id": template_id,
            "output_path": "",
            "warnings": ["structured_cv must be a dictionary."],
            "message": "Invalid structured CV.",
        }

    from services.ats_cv_postprocessing import clean_structured_cv_before_export
    structured_cv = clean_structured_cv_before_export(structured_cv)

    if not output_path:
        return {
            "success": False,
            "template_id": template_id,
            "output_path": "",
            "warnings": ["output_path is required."],
            "message": "Missing output path.",
        }

    try:
        destination = Path(output_path)
        destination.parent.mkdir(parents=True, exist_ok=True)

        document = Document()
        style = _style_for_template(template_id)
        _configure_document(document, style)
        warnings.extend(_render_header(document, structured_cv, style, photo_bytes, photo_filename))
        _render_sections(document, structured_cv, style)

        metadata = metadata if isinstance(metadata, dict) else {}
        if metadata:
            document.core_properties.comments = "Rendered by Job Application Assistant DOCX template service."

        document.save(str(destination))
        if not destination.exists() or destination.stat().st_size == 0:
            return {
                "success": False,
                "template_id": template_id,
                "output_path": "",
                "warnings": ["Rendered DOCX file was not created or is empty."],
                "message": "DOCX rendering failed.",
            }
        return {
            "success": True,
            "template_id": template_id,
            "output_path": str(destination),
            "warnings": warnings,
            "message": "DOCX rendered successfully.",
        }
    except Exception as exc:
        return {
            "success": False,
            "template_id": template_id,
            "output_path": "",
            "warnings": [str(exc)],
            "message": "DOCX rendering failed.",
        }


def _get_template(template_id: str) -> dict | None:
    for template in DOCX_TEMPLATE_CATALOG:
        if template["template_id"] == template_id:
            return deepcopy(template)
    return None


def _style_for_template(template_id: str) -> dict:
    if template_id == "ats_modern_docx":
        return {
            "margin": 0.75,
            "name_size": 20,
            "title_size": 11,
            "contact_size": 9,
            "body_size": 10,
            "heading_size": 12,
            "heading_space_before": 12,
            "heading_space_after": 3,
            "item_space_before": 6,
            "item_space_after": 2,
            "bullet_space_after": 1.5,
            "separator": True,
            "separator_color": "D0D0D0",
            "separator_size": "4",  # 1/2 pt thin line
            "align_header": WD_ALIGN_PARAGRAPH.CENTER,
            "modern_style": True,
            "accent_color": "303030",
            "heading_color": "303030",
            "bold_labels": True,
        }
    if template_id == "modern_professional":
        return {
            "margin": 0.58,
            "name_size": 25.5,
            "title_size": 11.8,
            "contact_size": 8.1,
            "body_size": 9.6,
            "heading_size": 10.7,
            "heading_space_before": 8.5,
            "heading_space_after": 1.8,
            "item_space_before": 5.2,
            "item_space_after": 1.1,
            "bullet_space_after": 0.5,
            "separator": True,
            "separator_color": "AEBAC6",
            "separator_size": "5",
            "align_header": WD_ALIGN_PARAGRAPH.CENTER,
            "modern_style": True,
            "accent_color": "274B63",
            "heading_color": "274B63",
            "bold_labels": True,
            "date_right_align": True,
            "date_tab_inch": 6.55,
            "header_rule_space_after": 10,
            "bullet_left_indent": 0.22,
            "bullet_first_line_indent": -0.13,
            "contact_max_chars": 104,
            "shorten_contact_links": True,
            "item_heading_color": "111827",
        }
    if template_id == "compact_technical":
        return {
            "margin": 0.48,
            "name_size": 19.5,
            "title_size": 10.2,
            "contact_size": 7.55,
            "body_size": 8.65,
            "heading_size": 9.7,
            "heading_space_before": 5.4,
            "heading_space_after": 0.8,
            "item_space_before": 2.9,
            "item_space_after": 0.25,
            "bullet_space_after": 0.05,
            "separator": True,
            "separator_color": "444444",
            "separator_size": "5",
            "align_header": WD_ALIGN_PARAGRAPH.LEFT,
            "modern_style": False,
            "accent_color": "111111",
            "heading_color": "111111",
            "bold_labels": True,
            "date_right_align": True,
            "date_tab_inch": 7.0,
            "header_rule_space_after": 6,
            "bullet_left_indent": 0.18,
            "bullet_first_line_indent": -0.11,
            "contact_max_chars": 112,
            "shorten_contact_links": True,
            "item_heading_color": "111111",
        }
    if template_id == "visual_photo_optional":
        return {
            "margin": 0.55,
            "name_size": 23.5,
            "title_size": 11.6,
            "contact_size": 8.0,
            "body_size": 9.35,
            "heading_size": 10.3,
            "heading_space_before": 7.4,
            "heading_space_after": 1.2,
            "item_space_before": 4.5,
            "item_space_after": 0.8,
            "bullet_space_after": 0.35,
            "separator": True,
            "separator_color": "8A8F98",
            "separator_size": "5",
            "align_header": WD_ALIGN_PARAGRAPH.LEFT,
            "modern_style": True,
            "accent_color": "24313D",
            "heading_color": "24313D",
            "bold_labels": True,
            "date_right_align": True,
            "date_tab_inch": 6.55,
            "supports_photo": True,
            "photo_width": 0.98,
            "photo_cell_width": 1.14,
            "header_rule_space_after": 9,
            "bullet_left_indent": 0.21,
            "bullet_first_line_indent": -0.12,
            "contact_max_chars": 98,
            "shorten_contact_links": True,
            "item_heading_color": "111827",
        }
    return {
        "margin": 0.75,
        "name_size": 16,
        "title_size": 10.5,
        "contact_size": 9,
        "body_size": 10,
        "heading_size": 11,
        "heading_space_before": 8,
        "heading_space_after": 2,
        "item_space_before": 4,
        "item_space_after": 1.5,
        "bullet_space_after": 0.8,
        "separator": True,
        "separator_color": "A0A0A0",
        "separator_size": "6",  # 3/4 pt line
        "align_header": WD_ALIGN_PARAGRAPH.LEFT,
        "modern_style": False,
        "accent_color": "222222",
        "heading_color": "222222",
        "bold_labels": True,
    }


def _configure_document(document: Document, style: dict) -> None:
    section = document.sections[0]
    margin = Inches(float(style["margin"]))
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(float(style["body_size"]))

    bullet_style = document.styles["List Bullet"]
    bullet_style.font.name = "Arial"
    bullet_style.font.size = Pt(float(style["body_size"]))
    bullet_style.paragraph_format.left_indent = Inches(float(style.get("bullet_left_indent", 0.25)))
    bullet_style.paragraph_format.first_line_indent = Inches(float(style.get("bullet_first_line_indent", -0.14)))


def _render_header(
    document: Document,
    structured_cv: dict,
    style: dict,
    photo_bytes: bytes | None = None,
    photo_filename: str = "",
) -> list[str]:
    warnings = []
    contact = _dict(structured_cv.get("contact"))
    full_name = _clean(contact.get("full_name"))
    target_title = _clean(contact.get("target_title"))
    contact_lines = _contact_lines(
        contact,
        int(style.get("contact_max_chars", 98)),
        shorten_links=bool(style.get("shorten_contact_links", False)),
    )

    if style.get("supports_photo") and photo_bytes:
        photo_warning = _render_photo_header(document, style, full_name, target_title, contact_lines, photo_bytes, photo_filename)
        if photo_warning:
            warnings.append(photo_warning)
        _add_header_rule(document, style)
        return warnings

    last_p = None

    if full_name:
        paragraph = document.add_paragraph()
        paragraph.alignment = style["align_header"]
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(full_name)
        run.bold = True
        run.font.size = Pt(float(style["name_size"]))
        run.font.color.rgb = _rgb(style.get("accent_color"))
        last_p = paragraph

    if target_title:
        paragraph = document.add_paragraph()
        paragraph.alignment = style["align_header"]
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(target_title)
        run.bold = True
        run.font.size = Pt(float(style["title_size"]))
        last_p = paragraph

    for index, contact_line in enumerate(contact_lines):
        paragraph = document.add_paragraph()
        paragraph.alignment = style["align_header"]
        paragraph.paragraph_format.space_after = Pt(1 if index < len(contact_lines) - 1 else 4)
        run = paragraph.add_run(contact_line)
        run.font.size = Pt(float(style["contact_size"]))
        run.font.color.rgb = _rgb("333333")
        last_p = paragraph

    if last_p and style.get("separator"):
        _add_bottom_border(
            last_p,
            color=style.get("separator_color", "CCCCCC"),
            size=style.get("separator_size", "6"),
            space="8"
        )
        last_p.paragraph_format.space_after = Pt(float(style.get("header_rule_space_after", 10)))
    return warnings


def _render_photo_header(
    document: Document,
    style: dict,
    full_name: str,
    target_title: str,
    contact_lines: list[str],
    photo_bytes: bytes,
    photo_filename: str,
) -> str:
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    photo_cell, text_cell = table.rows[0].cells
    photo_cell.width = Inches(float(style.get("photo_cell_width", 1.25)))
    text_cell.width = Inches(6.15)
    photo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    text_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    photo_paragraph = photo_cell.paragraphs[0]
    photo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    photo_paragraph.paragraph_format.space_before = Pt(1)
    try:
        run = photo_paragraph.add_run()
        run.add_picture(BytesIO(photo_bytes), width=Inches(float(style.get("photo_width", 1.08))))
    except Exception:
        photo_paragraph.add_run("")
        warning = "Photo could not be rendered; generated CV without the photo."
    else:
        warning = ""

    text_cell.paragraphs[0].text = ""
    if full_name:
        paragraph = text_cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(full_name)
        run.bold = True
        run.font.size = Pt(float(style["name_size"]))
        run.font.color.rgb = _rgb(style.get("accent_color"))
    if target_title:
        paragraph = text_cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(target_title)
        run.bold = True
        run.font.size = Pt(float(style["title_size"]))
        run.font.color.rgb = _rgb("333333")
    for index, contact_line in enumerate(contact_lines):
        paragraph = text_cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0 if index < len(contact_lines) - 1 else 2)
        run = paragraph.add_run(contact_line)
        run.font.size = Pt(float(style["contact_size"]))
        run.font.color.rgb = _rgb("333333")

    return warning


def _add_header_rule(document: Document, style: dict) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(float(style.get("header_rule_space_after", 10)))
    _add_bottom_border(
        paragraph,
        color=style.get("separator_color", "CCCCCC"),
        size=style.get("separator_size", "6"),
        space="4",
    )


def _render_sections(document: Document, structured_cv: dict, style: dict) -> None:
    for section_key in SECTION_ORDER:
        renderer = {
            "professional_summary": _summary_items,
            "skills": _skills_items,
            "experience": _experience_items,
            "projects": _project_items,
            "education": _education_items,
            "certifications": _certification_items,
            "languages": _language_items,
        }[section_key]

        items = renderer(structured_cv)

        # Robustness: Filter out items with empty content text to avoid rendering blank elements
        valid_items = []
        for item in items:
            cleaned_text = _clean(item.get("text"))
            if cleaned_text:
                valid_items.append({"type": item.get("type"), "text": cleaned_text})

        if not valid_items:
            continue

        _add_heading(document, SECTION_TITLES[section_key], style)
        for item in valid_items:
            if item["type"] == "heading":
                _add_item_heading(document, item["text"], style)
            elif item["type"] == "bullet":
                _add_bullet(document, item["text"], style)
            else:
                _add_paragraph(document, item["text"], style)


def _add_heading(document: Document, text: str, style: dict) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(float(style["heading_space_before"]))
    paragraph.paragraph_format.space_after = Pt(float(style["heading_space_after"]))
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(float(style["heading_size"]))
    run.font.color.rgb = _rgb(style.get("heading_color"))
    if style.get("separator"):
        _add_bottom_border(
            paragraph,
            color=style.get("separator_color", "CCCCCC"),
            size=style.get("separator_size", "6"),
            space="4"
        )


def _add_item_heading(document: Document, text: str, style: dict) -> None:
    if not text:
        return
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(float(style.get("item_space_before", 4)))
    paragraph.paragraph_format.space_after = Pt(float(style["item_space_after"]))

    parts = [p.strip() for p in text.split(" | ") if p.strip()]
    if not parts:
        return

    date_part = ""
    if style.get("date_right_align") and len(parts) >= 2 and _looks_like_date_range(parts[-1]):
        date_part = parts[-1]
        parts = parts[:-1]
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(float(style.get("date_tab_inch", 6.55))),
            WD_TAB_ALIGNMENT.RIGHT,
        )

    # First component (e.g. Job Title, Project Name, School Name) -> Bold
    run = paragraph.add_run(parts[0])
    run.bold = True
    run.font.size = Pt(float(style["body_size"]))
    run.font.color.rgb = _rgb(style.get("item_heading_color"))

    sep = "  •  " if style.get("modern_style") else "  |  "

    for i, part in enumerate(parts[1:]):
        sep_run = paragraph.add_run(sep)
        sep_run.font.size = Pt(float(style["body_size"]))

        run = paragraph.add_run(part)
        run.font.size = Pt(float(style["body_size"]))

        # Second component (e.g. Company name, Degree type) -> Italic
        if i == 0:
            run.italic = True

    if date_part:
        date_run = paragraph.add_run("\t" + date_part)
        date_run.font.size = Pt(float(style["body_size"]))


def _add_paragraph(document: Document, text: str, style: dict) -> None:
    if not text:
        return
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(float(style["item_space_after"]))
    if style.get("bold_labels") and ":" in text:
        label, rest = text.split(":", 1)
        if 1 <= len(label.split()) <= 4 and len(label) <= 32:
            label_run = paragraph.add_run(f"{label}:")
            label_run.bold = True
            label_run.font.size = Pt(float(style["body_size"]))
            rest_run = paragraph.add_run(rest)
            rest_run.font.size = Pt(float(style["body_size"]))
            return
    run = paragraph.add_run(text)
    run.font.size = Pt(float(style["body_size"]))


def _add_bullet(document: Document, text: str, style: dict) -> None:
    if not text:
        return
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(float(style["bullet_space_after"]))
    paragraph.paragraph_format.left_indent = Inches(float(style.get("bullet_left_indent", 0.25)))
    paragraph.paragraph_format.first_line_indent = Inches(float(style.get("bullet_first_line_indent", -0.14)))
    run = paragraph.add_run(text)
    run.font.size = Pt(float(style["body_size"]))


def _summary_items(structured_cv: dict) -> list[dict]:
    summary = _clean(
        structured_cv.get("professional_summary")
        or structured_cv.get("summary")
        or structured_cv.get("career_objective")
        or structured_cv.get("technical_summary")
    )
    return [{"type": "paragraph", "text": summary}] if summary else []


def _skills_items(structured_cv: dict) -> list[dict]:
    skills = structured_cv.get("skills") or structured_cv.get("technical_skills") or structured_cv.get("core_skills")
    items = []
    if isinstance(skills, dict):
        for category, values in skills.items():
            flattened = _flatten(values)
            if flattened:
                items.append({"type": "paragraph", "text": f"{_label(category)}: {', '.join(flattened)}"})
    else:
        flattened = _flatten(skills)
        if flattened:
            items.append({"type": "paragraph", "text": ", ".join(flattened)})
    return items


def _experience_items(structured_cv: dict) -> list[dict]:
    records = _list(structured_cv.get("experience")) + _list(structured_cv.get("internship_experience"))
    items = []
    for record in records:
        if not isinstance(record, dict):
            continue
        heading = _join_non_empty([
            record.get("title") or record.get("role") or record.get("position"),
            record.get("company") or record.get("organization"),
            record.get("location"),
            _date_range(record),
        ], " | ")
        if heading:
            items.append({"type": "heading", "text": heading})
        for bullet in _list(record.get("bullets")):
            cleaned_bullet = _clean(bullet)
            if cleaned_bullet:
                items.append({"type": "bullet", "text": cleaned_bullet})
        description = _clean(record.get("description"))
        if description:
            items.append({"type": "paragraph", "text": description})
    return items


def _project_items(structured_cv: dict) -> list[dict]:
    items = []
    for record in _list(structured_cv.get("projects")):
        if not isinstance(record, dict):
            continue
        heading = _join_non_empty([record.get("name") or record.get("title"), _date_range(record)], " | ")
        if heading:
            items.append({"type": "heading", "text": heading})
        technologies = _flatten(record.get("technologies"))
        if technologies:
            items.append({"type": "paragraph", "text": f"Technologies: {', '.join(technologies)}"})
        description = _clean(record.get("description"))
        if description:
            items.append({"type": "paragraph", "text": description})
        for bullet in _list(record.get("bullets")):
            cleaned_bullet = _clean(bullet)
            if cleaned_bullet:
                items.append({"type": "bullet", "text": cleaned_bullet})
    return items


def _education_items(structured_cv: dict) -> list[dict]:
    items = []
    for record in _list(structured_cv.get("education")):
        if not isinstance(record, dict):
            continue
        heading = _join_non_empty([
            record.get("school") or record.get("institution") or record.get("university"),
            record.get("degree"),
            record.get("department") or record.get("field"),
            _date_range(record),
        ], " | ")
        if heading:
            items.append({"type": "heading", "text": heading})
        for detail in _list(record.get("details")):
            cleaned_detail = _clean(detail)
            if cleaned_detail:
                items.append({"type": "bullet", "text": cleaned_detail})
    return items


def _certification_items(structured_cv: dict) -> list[dict]:
    items = []
    for record in _list(structured_cv.get("certifications")):
        if not isinstance(record, dict):
            continue
        name = record.get("name") or record.get("certification")
        issuer = record.get("issuer") or record.get("organization")
        display_issuer = issuer
        if name and issuer:
            name_cleaned = str(name).strip()
            norm_issuer = str(issuer).strip().lower()
            import re
            name_stripped = re.sub(r'[\s\-_—|~]+$', '', name_cleaned)
            try:
                pattern = r'\b' + re.escape(norm_issuer) + r'$'
                if re.search(pattern, name_stripped.lower()):
                    display_issuer = None
            except Exception:
                if name_stripped.lower().endswith(norm_issuer):
                    display_issuer = None
        line = _join_non_empty([
            name,
            display_issuer,
            record.get("date"),
            record.get("link"),
        ], " | ")
        if line:
            items.append({"type": "paragraph", "text": line})
    return items


def _language_items(structured_cv: dict) -> list[dict]:
    items = []
    for record in _list(structured_cv.get("languages")):
        if isinstance(record, dict):
            line = _join_non_empty([record.get("language"), record.get("level")], " - ")
        else:
            line = _clean(record)
        if line:
            items.append({"type": "paragraph", "text": line})
    return items


def _contact_lines(contact: dict, max_chars: int, shorten_links: bool = False) -> list[str]:
    items = [
        (key, _display_contact_value(key, contact.get(key), shorten_links))
        for key in ["email", "phone", "location", "linkedin", "github", "portfolio"]
        if _display_contact_value(key, contact.get(key), shorten_links)
    ]
    if not items:
        return []

    one_line = _join_non_empty([value for _, value in items], " | ")
    if len(one_line) <= max_chars:
        return [one_line]

    primary = [value for key, value in items if key not in {"linkedin", "github", "portfolio"}]
    links = [(key, value) for key, value in items if key in {"linkedin", "github", "portfolio"}]
    lines = []
    if primary:
        lines.append(_join_non_empty(primary, " | "))

    link_values = [value for _, value in links]
    if len(link_values) <= 2:
        if link_values:
            lines.append(_join_non_empty(link_values, " | "))
        return lines

    social_values = [value for key, value in links if key in {"linkedin", "github"}]
    portfolio_values = [value for key, value in links if key == "portfolio"]
    if social_values:
        lines.append(_join_non_empty(social_values, " | "))
    lines.extend(portfolio_values)
    return [line for line in lines if line]


def _display_contact_value(key: str, value: Any, shorten_links: bool = False) -> str:
    cleaned = _clean(value)
    if shorten_links and key in {"linkedin", "github", "portfolio"}:
        cleaned = cleaned.rstrip("/")
        for prefix in ("https://www.", "http://www.", "https://", "http://", "www."):
            if cleaned.lower().startswith(prefix):
                return cleaned[len(prefix):]
    return cleaned


def _date_range(record: dict) -> str:
    return _join_non_empty([
        record.get("start_date") or record.get("start"),
        record.get("end_date") or record.get("end"),
    ], " - ")


def _join_non_empty(values: list[Any], separator: str) -> str:
    return separator.join(_clean(value) for value in values if _clean(value))


def _clean(value: Any) -> str:
    return str(value or "").strip()


def _dict(value: Any) -> dict:
    return value if isinstance(value, dict) else {}


def _list(value: Any) -> list:
    return value if isinstance(value, list) else []


def _flatten(value: Any) -> list[str]:
    if isinstance(value, list):
        return [_clean(item) for item in value if _clean(item)]
    if isinstance(value, dict):
        flattened = []
        for nested in value.values():
            flattened.extend(_flatten(nested))
        return flattened
    if isinstance(value, str):
        return [part.strip() for part in value.replace(";", ",").split(",") if part.strip()]
    return []


def _label(value: Any) -> str:
    return _clean(value).replace("_", " ").title()


def _rgb(value: str | None) -> RGBColor:
    text = str(value or "222222").strip().lstrip("#")
    if not re_fullmatch_hex(text):
        text = "222222"
    return RGBColor(int(text[0:2], 16), int(text[2:4], 16), int(text[4:6], 16))


def re_fullmatch_hex(text: str) -> bool:
    if len(text) != 6:
        return False
    return all(character in "0123456789abcdefABCDEF" for character in text)


def _looks_like_date_range(value: str) -> bool:
    text = _clean(value).lower()
    if not text:
        return False
    date_terms = {
        "present", "current", "now", "bugün", "bugun", "devam", "ocak", "şubat", "subat",
        "mart", "nisan", "mayıs", "mayis", "haziran", "temmuz", "ağustos", "agustos",
        "eylül", "eylul", "ekim", "kasım", "kasim", "aralık", "aralik",
    }
    if any(term in text for term in date_terms):
        return True
    return any(char.isdigit() for char in text) and any(sep in text for sep in ("-", "–", "/", "."))
