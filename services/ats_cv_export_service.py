import re
from io import BytesIO
from copy import deepcopy
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from fastapi import HTTPException
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import HRFlowable, Image as RLImage, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from services.ats_cv_relevance import (
    rank_certifications_for_job,
    rank_education_for_job,
    rank_experience_for_job,
    rank_projects_for_job,
    rank_skills_for_job,
    relevance_keywords,
)
from services.cv_photo_service import prepare_cv_photo_for_export


SECTION_TITLES = {
    "English": {
        "professional_summary": "Professional Summary",
        "career_objective": "Career Objective",
        "technical_summary": "Technical Summary",
        "skills": "Skills",
        "core_skills": "Core Skills",
        "technical_skills": "Technical Skills",
        "experience": "Experience",
        "internship_experience": "Internship Experience",
        "projects": "Projects",
        "education": "Education",
        "certifications": "Certifications",
        "languages": "Languages",
    },
    "Turkish": {
        "professional_summary": "Profesyonel Özet",
        "career_objective": "Kariyer Hedefi",
        "technical_summary": "Teknik Özet",
        "skills": "Yetenekler",
        "core_skills": "Temel Yetkinlikler",
        "technical_skills": "Teknik Yetenekler",
        "experience": "Deneyim",
        "internship_experience": "Staj Deneyimi",
        "projects": "Projeler",
        "education": "Eğitim",
        "certifications": "Sertifikalar",
        "languages": "Diller",
    },
}

PDF_FONT_NAME = "Helvetica"
PDF_BOLD_FONT_NAME = "Helvetica-Bold"

PRESERVED_CONTACT_FIELDS = [
    "full_name",
    "email",
    "phone",
    "linkedin",
    "github",
    "portfolio",
]

PRESERVED_RECORD_FIELDS = {
    "experience": ["company"],
    "projects": ["name"],
    "education": ["school"],
    "certifications": ["name", "issuer"],
}

def render_ats_cv_to_docx(
    ats_cv: dict,
    template: dict,
    language: str,
    one_page: bool = False,
    enabled_sections: set[str] | None = None,
    export_style: str = "standard",
    photo_bytes: bytes | None = None,
    photo_filename: str = "",
) -> bytes:
    try:
        export_style = _effective_export_style(export_style, one_page)
        render_cv = _cv_for_export_style(ats_cv, template, language, export_style)
        document = Document()
        template_style = _template_style(template, export_style, estimate_cv_content_density(render_cv))
        _configure_docx(document, template_style)

        contact = render_contact(
            render_cv,
            language,
            int(template_style.get("contact_max_chars", 104)),
            shorten_links=bool(template_style.get("shorten_contact_links", False)),
            separator=template_style.get("contact_separator", " | "),
        )
        if _section_enabled("contact", enabled_sections):
            _add_docx_header(document, contact, template_style, photo_bytes, photo_filename)

        for section in _ordered_sections(render_cv, template, language, enabled_sections):
            _add_docx_section(document, section, template_style, language)

        output = BytesIO()
        document.save(output)
        return output.getvalue()
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(exc)}") from exc


def render_ats_cv_to_pdf(
    ats_cv: dict,
    template: dict,
    language: str,
    one_page: bool = False,
    enabled_sections: set[str] | None = None,
    export_style: str = "standard",
    photo_bytes: bytes | None = None,
    photo_filename: str = "",
) -> bytes:
    try:
        export_style = _effective_export_style(export_style, one_page)
        render_cv = _cv_for_export_style(ats_cv, template, language, export_style)
        _register_pdf_fonts()
        template_style = _template_style(template, export_style, estimate_cv_content_density(render_cv))
        output = BytesIO()
        document = SimpleDocTemplate(
            output,
            pagesize=A4,
            rightMargin=template_style["pdf_margin"] * inch,
            leftMargin=template_style["pdf_margin"] * inch,
            topMargin=template_style["pdf_margin"] * inch,
            bottomMargin=template_style["pdf_margin"] * inch,
        )
        styles = _pdf_styles(template_style)
        story = []

        contact = render_contact(
            render_cv,
            language,
            int(template_style.get("contact_max_chars", 104)),
            shorten_links=bool(template_style.get("shorten_contact_links", False)),
            separator=template_style.get("contact_separator", " | "),
        )
        if _section_enabled("contact", enabled_sections):
            story.extend(_pdf_header(contact, styles, template_style, photo_bytes, photo_filename))

        for section in _ordered_sections(render_cv, template, language, enabled_sections):
            story.append(Spacer(1, template_style["pdf_section_spacing"]))
            story.append(Paragraph(escape(_heading_text(section["heading"], language)), styles["SectionHeading"]))
            if template_style["heading_separator"]:
                story.append(_pdf_horizontal_line(template_style))
            for item in section["items"]:
                if item["type"] == "paragraph":
                    story.append(Paragraph(escape(item["text"]), styles["Body"]))
                elif item["type"] == "heading":
                    story.append(Paragraph(_pdf_item_heading_markup(item["text"], template_style), styles["ItemHeading"]))
                elif item["type"] == "bullet":
                    story.append(Paragraph(f"- {escape(item['text'])}", styles["Bullet"]))
                elif item["type"] == "separator":
                    story.append(Paragraph(escape(item["text"]), styles["Separator"]))
            story.append(Spacer(1, template_style["pdf_section_after_spacing"]))

        document.build(story)
        return output.getvalue()
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(exc)}") from exc


def build_plain_text_preview(
    ats_cv: dict,
    template: dict,
    language: str,
    one_page: bool = False,
    enabled_sections: set[str] | None = None,
    export_style: str = "standard",
) -> str:
    export_style = _effective_export_style(export_style, one_page)
    render_cv = _cv_for_export_style(ats_cv, template, language, export_style)
    template_style = _template_style(template, export_style, estimate_cv_content_density(render_cv))
    contact = render_contact(
        render_cv,
        language,
        int(template_style.get("contact_max_chars", 104)),
        shorten_links=False,
        separator=" | ",
    )
    lines = []

    if _section_enabled("contact", enabled_sections) and contact["full_name"]:
        lines.append(contact["full_name"])
    if _section_enabled("contact", enabled_sections) and contact["target_title"]:
        lines.append(contact["target_title"])
    if _section_enabled("contact", enabled_sections):
        lines.extend(contact["contact_lines"])

    for section in _ordered_sections(render_cv, template, language, enabled_sections):
        if lines:
            lines.append("")
        lines.append(_heading_text(section["heading"], language))
        if template_style["heading_separator"] and not template_style.get("separator_as_rule"):
            lines.append("-" * max(12, len(_heading_text(section["heading"], language))))
        for item in section["items"]:
            if item["type"] == "bullet":
                lines.append(f"- {item['text']}")
            else:
                lines.append(item["text"])

    return "\n".join(lines).strip()


def compact_ats_cv_for_one_page(ats_cv: dict, template: dict, language: str) -> dict:
    return balance_one_page_content(ats_cv, template, language)


def balance_one_page_content(ats_cv: dict, template: dict, language: str) -> dict:
    compact_cv = deepcopy(ats_cv)
    density = estimate_cv_content_density(compact_cv)
    limits = _density_limits(density)
    metadata = compact_cv.get("ats_metadata", {}) if isinstance(compact_cv, dict) else {}
    keywords = relevance_keywords("", metadata)

    for summary_key in ["professional_summary", "career_objective", "technical_summary"]:
        compact_cv[summary_key] = _trim_text(compact_cv.get(summary_key), limits["summary_chars"])

    skills = compact_cv.get("skills", {})
    if isinstance(skills, dict):
        ranked_skills = rank_skills_for_job(skills, "", metadata)
        for group, limit in limits["skills"].items():
            skills[group] = _prioritize_strings(_clean_list(ranked_skills.get(group, [])), keywords, limit)

    compact_cv["experience"] = rank_experience_for_job(compact_cv.get("experience", []), "", metadata)
    compact_cv["projects"] = rank_projects_for_job(compact_cv.get("projects", []), "", metadata)
    compact_cv["education"] = rank_education_for_job(compact_cv.get("education", []), "", metadata)
    compact_cv["certifications"] = rank_certifications_for_job(compact_cv.get("certifications", []), "", metadata)

    for record in compact_cv.get("experience", []):
        if isinstance(record, dict):
            bullets = _clean_list(record.get("bullets", []))
            record["bullets"] = _prioritize_strings(bullets, keywords, limits["experience_bullets"]) or bullets[:1]

    for record in compact_cv.get("projects", []):
        if isinstance(record, dict):
            record["description"] = _trim_text(record.get("description"), limits["project_description_chars"])
            bullets = _clean_list(record.get("bullets", []))
            record["bullets"] = _prioritize_strings(bullets, keywords, limits["project_bullets"]) or bullets[:1]
            record["technologies"] = _prioritize_strings(
                _clean_list(record.get("technologies", [])),
                keywords,
                limits["project_technologies"],
            )

    for record in compact_cv.get("education", []):
        if isinstance(record, dict):
            details = _clean_list(record.get("details", []))
            record["details"] = _prioritize_strings(details, keywords, limits["education_details"]) or details[:1]

    compact_cv["certifications"] = [
        certification for certification in compact_cv.get("certifications", [])
        if isinstance(certification, dict) and any(_clean_text(value) for value in certification.values())
    ][:limits["certifications"]]

    return compact_cv


def estimate_cv_content_density(ats_cv: dict) -> str:
    text_units = []
    for key in ["professional_summary", "career_objective", "technical_summary"]:
        text_units.append(_clean_text(ats_cv.get(key)))

    skills = ats_cv.get("skills", {})
    if isinstance(skills, dict):
        for values in skills.values():
            text_units.extend(_clean_list(values))

    bullet_count = 0
    record_count = 0
    for section_key in ["experience", "projects", "education", "certifications", "languages"]:
        records = ats_cv.get(section_key, [])
        if not isinstance(records, list):
            continue
        for record in records:
            if not isinstance(record, dict):
                continue
            record_count += 1
            for value in record.values():
                if isinstance(value, list):
                    bullet_count += len(value)
                    text_units.extend(_clean_list(value))
                else:
                    text_units.append(_clean_text(value))

    char_count = sum(len(text) for text in text_units if text)
    density_score = char_count + (bullet_count * 90) + (record_count * 120)
    if density_score < 4300:
        return "short"
    if density_score < 7200:
        return "medium"
    return "long"


def render_contact(
    ats_cv: dict,
    language: str,
    max_chars: int = 104,
    shorten_links: bool = True,
    separator: str = " | ",
) -> dict:
    from services.ats_cv_postprocessing import _clean_character_spacing
    contact = ats_cv.get("contact", {}) if isinstance(ats_cv, dict) else {}
    contact_items = []
    
    full_name = _clean_character_spacing(contact.get("full_name", ""))
    target_title = _clean_character_spacing(contact.get("target_title", ""))
    
    seen_links = set()
    for key in ["email", "phone", "location", "linkedin", "github", "portfolio"]:
        val = contact.get(key, "")
        val = _clean_character_spacing(val)
        if not val:
            continue
            
        if key in ["linkedin", "github", "portfolio"]:
            # Normalize to detect duplicates (e.g. protocol, www, trailing slash)
            norm = val.lower().strip().rstrip("/")
            norm = re.sub(r"^https?://(www\.)?", "", norm)
            if norm in seen_links:
                continue
            seen_links.add(norm)
        display_value = _display_contact_value(key, val) if shorten_links else val
        if display_value:
            contact_items.append((key, display_value))
        
    contact_lines = _contact_lines(contact_items, max_chars, separator)

    return {
        "full_name": full_name,
        "target_title": target_title,
        "contact_line": separator.join(value for _, value in contact_items),
        "contact_lines": contact_lines,
    }


def _contact_lines(contact_items: list[tuple[str, str]], max_chars: int = 104, separator: str = " | ") -> list[str]:
    """Pack contact values without splitting URLs; keep LinkedIn/GitHub together when wrapping."""
    if not contact_items:
        return []

    one_line = separator.join(value for _, value in contact_items)
    if len(one_line) <= max_chars:
        return [one_line]

    primary = [value for key, value in contact_items if key not in {"linkedin", "github", "portfolio"}]
    links = [(key, value) for key, value in contact_items if key in {"linkedin", "github", "portfolio"}]
    lines = []
    if primary:
        lines.append(separator.join(primary))

    social_values = [value for key, value in links if key in {"linkedin", "github"}]
    portfolio_values = [value for key, value in links if key == "portfolio"]
    if social_values:
        lines.append(separator.join(social_values))
    lines.extend(portfolio_values)

    return [line for line in lines if line]


def _display_contact_value(key: str, value: str) -> str:
    cleaned = _clean_text(value)
    if key in {"linkedin", "github", "portfolio"}:
        cleaned = cleaned.rstrip("/")
        for prefix in ("https://www.", "http://www.", "https://", "http://", "www."):
            if cleaned.lower().startswith(prefix):
                return cleaned[len(prefix):]
    return cleaned


def render_summary(ats_cv: dict, language: str, section_key: str = "professional_summary") -> dict | None:
    value = _clean_text(ats_cv.get(section_key))
    if not value:
        return None
    return {
        "key": section_key,
        "heading": _section_title(section_key, language),
        "items": [{"type": "paragraph", "text": value}],
    }


def render_skills(ats_cv: dict, language: str, section_key: str = "skills") -> dict | None:
    skills = ats_cv.get("skills", {})
    if not isinstance(skills, dict):
        return None

    items = []
    if section_key in {"technical_skills", "core_skills"}:
        skill_groups = [section_key]
    else:
        skill_groups = ["technical_skills", "core_skills", "tools", "databases", "cloud", "soft_skills"]

    single_group_section = section_key in {"technical_skills", "core_skills"}

    for group in skill_groups:
        if section_key == "core_skills" and group == "core_skills":
            values = _modern_clean_core_skill_values(skills, ats_cv.get("ats_metadata", {}))
        else:
            values = _clean_list(skills.get(group, []))
        if values:
            text = ", ".join(values) if single_group_section else f"{_skill_group_title(group, language)}: {', '.join(values)}"
            items.append({
                "type": "paragraph",
                "text": text,
            })

    if not items:
        return None

    return {
        "key": "skills",
        "heading": _section_title(section_key, language),
        "items": items,
    }


def render_experience(ats_cv: dict, language: str, section_key: str = "experience") -> dict | None:
    records = ats_cv.get("experience", [])
    if not isinstance(records, list):
        return None

    items = []
    for record in records:
        if not isinstance(record, dict):
            continue
        heading = _join_non_empty([
            record.get("role"),
            record.get("company"),
            record.get("location"),
            _date_range(record.get("start_date"), record.get("end_date")),
        ])
        bullets = _clean_list(record.get("bullets", []))
        if heading:
            items.append({"type": "heading", "text": heading})
        for bullet in bullets:
            items.append({"type": "bullet", "text": bullet})

    if not items:
        return None

    return {
        "key": "experience",
        "heading": _section_title(section_key, language),
        "items": items,
    }


def render_projects(ats_cv: dict, language: str, template: dict | None = None) -> dict | None:
    records = ats_cv.get("projects", [])
    if not isinstance(records, list):
        return None

    items = []
    template_id = (template or {}).get("id", "")
    contact = ats_cv.get("contact", {}) if isinstance(ats_cv, dict) else {}
    contact_links = {
        _normalize_url_for_compare(contact.get("linkedin")),
        _normalize_url_for_compare(contact.get("github")),
        _normalize_url_for_compare(contact.get("portfolio")),
    }
    contact_links.discard("")
    rendered_count = 0
    for record in records:
        if not isinstance(record, dict):
            continue
        name = _clean_text(record.get("name"))
        description = _clean_text(record.get("description"))
        technologies = _clean_list(record.get("technologies", []))
        bullets = _clean_list(record.get("bullets", []))
        link = _clean_text(record.get("link"))

        if name:
            items.append({"type": "heading", "text": name})
        if description:
            items.append({"type": "paragraph", "text": description})
        if technologies:
            label = "Technologies" if _language_key(language) == "English" else "Teknolojiler"
            items.append({"type": "paragraph", "text": f"{label}: {', '.join(technologies)}"})
        for bullet in bullets:
            items.append({"type": "bullet", "text": bullet})
        link_label = _project_link_label(link, contact_links, language)
        if link_label:
            items.append({"type": "paragraph", "text": link_label})
        rendered_count += 1
        if template_id == "technical_developer" and rendered_count < len(records):
            items.append({"type": "separator", "text": "-" * 32})

    if not items:
        return None

    return {
        "key": "projects",
        "heading": _section_title("projects", language),
        "items": items,
    }


def render_education(ats_cv: dict, language: str) -> dict | None:
    records = ats_cv.get("education", [])
    if not isinstance(records, list):
        return None

    items = []
    for record in records:
        if not isinstance(record, dict):
            continue
        heading = _join_non_empty([
            record.get("school"),
            record.get("degree"),
            record.get("department"),
            _date_range(record.get("start_date"), record.get("end_date")),
        ])
        details = _clean_list(record.get("details", []))
        if heading:
            items.append({"type": "heading", "text": heading})
        for detail in details:
            items.append({"type": "bullet", "text": detail})

    if not items:
        return None

    return {
        "key": "education",
        "heading": _section_title("education", language),
        "items": items,
    }


def render_certifications(ats_cv: dict, language: str) -> dict | None:
    records = ats_cv.get("certifications", [])
    if not isinstance(records, list):
        return None

    items = []
    for record in records:
        if not isinstance(record, dict):
            continue
        name = record.get("name")
        issuer = record.get("issuer")
        display_issuer = issuer
        if name and issuer:
            name_cleaned = str(name).strip()
            norm_issuer = str(issuer).strip().lower()
            name_stripped = re.sub(r'[\s\-_—|~]+$', '', name_cleaned)
            try:
                pattern = r'\b' + re.escape(norm_issuer) + r'$'
                if re.search(pattern, name_stripped.lower()):
                    display_issuer = None
            except Exception:
                if name_stripped.lower().endswith(norm_issuer):
                    display_issuer = None
        line = _join_non_empty([
            name,
            display_issuer,
            record.get("date"),
            record.get("link"),
        ])
        if line:
            items.append({"type": "bullet", "text": line})

    if not items:
        return None

    return {
        "key": "certifications",
        "heading": _section_title("certifications", language),
        "items": items,
    }


def render_languages(ats_cv: dict, language: str) -> dict | None:
    records = ats_cv.get("languages", [])
    if not isinstance(records, list):
        return None

    items = []
    for record in records:
        if not isinstance(record, dict):
            continue
        line = _join_non_empty([record.get("language"), record.get("level")], separator=" - ")
        if line:
            items.append({"type": "bullet", "text": line})

    if not items:
        return None

    return {
        "key": "languages",
        "heading": _section_title("languages", language),
        "items": items,
    }


def _ordered_sections(
    ats_cv: dict,
    template: dict,
    language: str,
    enabled_sections: set[str] | None = None,
) -> list[dict]:
    sections = []
    used_keys = set()

    for section_key in template.get("section_order", []):
        canonical_section = _canonical_section(section_key)
        if not _section_enabled(canonical_section, enabled_sections):
            continue
        section = _render_section_by_key(ats_cv, language, section_key, template)
        if section and section["key"] not in used_keys:
            sections.append(section)
            used_keys.add(section["key"])

    return sections


def _canonical_section(section_key: str) -> str:
    if section_key in {"professional_summary", "career_objective", "technical_summary"}:
        return "summary"
    if section_key in {"core_skills", "technical_skills"}:
        return "skills"
    if section_key == "internship_experience":
        return "experience"
    return section_key


def _section_enabled(section_key: str, enabled_sections: set[str] | None) -> bool:
    if enabled_sections is None:
        return True
    return section_key in enabled_sections


def _render_section_by_key(ats_cv: dict, language: str, section_key: str, template: dict | None = None) -> dict | None:
    if section_key in {"contact", "title"}:
        return None
    if section_key in {"professional_summary", "career_objective", "technical_summary"}:
        return render_summary(ats_cv, language, section_key)
    if section_key in {"skills", "core_skills", "technical_skills"}:
        return render_skills(ats_cv, language, section_key)
    if section_key in {"experience", "internship_experience"}:
        return render_experience(ats_cv, language, section_key)
    if section_key == "projects":
        return render_projects(ats_cv, language, template)
    if section_key == "education":
        return render_education(ats_cv, language)
    if section_key == "certifications":
        return render_certifications(ats_cv, language)
    if section_key == "languages":
        return render_languages(ats_cv, language)
    return None


def _configure_docx(document: Document, template_style: dict) -> None:
    section = document.sections[0]
    section.top_margin = Inches(template_style["docx_margin"])
    section.bottom_margin = Inches(template_style["docx_margin"])
    section.left_margin = Inches(template_style["docx_margin"])
    section.right_margin = Inches(template_style["docx_margin"])

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(template_style["docx_body_size"])

    bullet_style = document.styles["List Bullet"]
    bullet_style.font.name = "Arial"
    bullet_style.font.size = Pt(template_style["docx_body_size"])
    bullet_style.paragraph_format.left_indent = Inches(float(template_style.get("docx_bullet_left_indent", 0.24)))
    bullet_style.paragraph_format.first_line_indent = Inches(float(template_style.get("docx_bullet_first_line_indent", -0.13)))


def _add_docx_header(
    document: Document,
    contact: dict,
    template_style: dict,
    photo_bytes: bytes | None = None,
    photo_filename: str = "",
) -> None:
    if template_style.get("supports_photo") and photo_bytes:
        if _add_docx_photo_header(document, contact, template_style, photo_bytes, photo_filename):
            _add_docx_header_rule(document, template_style)
            return

    alignment = _docx_header_alignment(template_style)
    last_paragraph = None
    if contact["full_name"]:
        paragraph = document.add_paragraph()
        paragraph.alignment = alignment
        paragraph.paragraph_format.space_after = Pt(1.5)
        run = paragraph.add_run(contact["full_name"])
        run.bold = True
        run.font.size = Pt(template_style["docx_name_size"])
        run.font.color.rgb = _docx_rgb(template_style.get("accent_color", "111111"))
        last_paragraph = paragraph

    if contact["target_title"]:
        paragraph = document.add_paragraph()
        paragraph.alignment = alignment
        paragraph.paragraph_format.space_after = Pt(1.0)
        run = paragraph.add_run(contact["target_title"])
        run.bold = True
        run.italic = bool(template_style.get("italic_title", False))
        run.font.size = Pt(template_style["docx_title_size"])
        run.font.color.rgb = _docx_rgb(template_style.get("title_color", "111111"))
        last_paragraph = paragraph

    for index, contact_line in enumerate(contact["contact_lines"]):
        paragraph = document.add_paragraph()
        paragraph.alignment = alignment
        paragraph.paragraph_format.space_after = Pt(0.4 if index < len(contact["contact_lines"]) - 1 else 3.0)
        run = paragraph.add_run(contact_line)
        run.font.size = Pt(template_style["docx_contact_size"])
        run.font.color.rgb = _docx_rgb("333333")
        last_paragraph = paragraph

    if last_paragraph and template_style.get("header_separator"):
        _add_docx_header_rule(document, template_style)


def _add_docx_photo_header(
    document: Document,
    contact: dict,
    template_style: dict,
    photo_bytes: bytes,
    photo_filename: str,
) -> bool:
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    photo_cell, text_cell = table.rows[0].cells
    photo_cell.width = Inches(float(template_style.get("docx_photo_cell_width", 1.08)))
    text_cell.width = Inches(6.1)
    photo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    text_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    photo_paragraph = photo_cell.paragraphs[0]
    photo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    photo_paragraph.paragraph_format.space_before = Pt(1)
    photo_paragraph.paragraph_format.space_after = Pt(1)
    try:
        photo_bytes = prepare_cv_photo_for_export(photo_bytes)
        run = photo_paragraph.add_run()
        run.add_picture(BytesIO(photo_bytes), width=Inches(float(template_style.get("docx_photo_width", 0.92))))
    except Exception:
        table._element.getparent().remove(table._element)
        return False

    text_cell.paragraphs[0].text = ""
    if contact["full_name"]:
        paragraph = text_cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(1.0)
        run = paragraph.add_run(contact["full_name"])
        run.bold = True
        run.font.size = Pt(template_style["docx_name_size"])
        run.font.color.rgb = _docx_rgb(template_style.get("accent_color", "111111"))
    if contact["target_title"]:
        paragraph = text_cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0.8)
        run = paragraph.add_run(contact["target_title"])
        run.bold = True
        run.italic = bool(template_style.get("italic_title", False))
        run.font.size = Pt(template_style["docx_title_size"])
        run.font.color.rgb = _docx_rgb(template_style.get("title_color", "111111"))
    for index, contact_line in enumerate(contact["contact_lines"]):
        paragraph = text_cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0.2 if index < len(contact["contact_lines"]) - 1 else 1.8)
        run = paragraph.add_run(contact_line)
        run.font.size = Pt(template_style["docx_contact_size"])
        run.font.color.rgb = _docx_rgb("333333")

    return True


def _add_docx_header_rule(document: Document, template_style: dict) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(float(template_style.get("docx_header_rule_space_after", 6)))
    _add_docx_bottom_border(paragraph, template_style.get("separator_color", "B7B7B7"), "5", "2")


def _add_docx_section(document: Document, section: dict, template_style: dict, language: str) -> None:
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(template_style["docx_section_space_before"])
    heading.paragraph_format.space_after = Pt(template_style["docx_heading_space_after"])
    run = heading.add_run(_heading_text(section["heading"], language))
    run.bold = True
    run.font.size = Pt(template_style["docx_heading_size"])
    run.font.color.rgb = _docx_rgb(template_style.get("heading_color", "111111"))

    if template_style["heading_separator"]:
        if template_style.get("separator_as_rule"):
            _add_docx_horizontal_line(document, template_style)
        else:
            separator = document.add_paragraph()
            separator.paragraph_format.space_after = Pt(template_style["docx_separator_space_after"])
            run = separator.add_run(template_style["separator_text"])
            run.font.size = Pt(template_style["docx_separator_size"])

    for item in section["items"]:
        if item["type"] == "heading":
            _add_docx_item_heading(document, item["text"], template_style)
        elif item["type"] == "bullet":
            _add_docx_bullet(document, item["text"], template_style)
        elif item["type"] == "separator":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(template_style["docx_item_space_after"])
            paragraph.add_run(item["text"])
        else:
            _add_docx_body_paragraph(document, item["text"], template_style)


def _add_docx_item_heading(document: Document, text: str, template_style: dict) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(float(template_style.get("docx_item_space_before", 1.8)))
    paragraph.paragraph_format.space_after = Pt(template_style["docx_item_space_after"])

    parts = [part.strip() for part in text.split(" | ") if part.strip()]
    if not parts:
        return

    date_part = ""
    if len(parts) >= 2 and _looks_like_date_range(parts[-1]):
        date_part = parts[-1]
        parts = parts[:-1]
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(float(template_style.get("docx_date_tab_inch", 6.45))),
            WD_TAB_ALIGNMENT.RIGHT,
        )

    first_run = paragraph.add_run(parts[0])
    first_run.bold = True
    first_run.font.size = Pt(template_style["docx_body_size"])
    first_run.font.color.rgb = _docx_rgb(template_style.get("item_heading_color", "111111"))

    separator = "  •  " if template_style.get("item_separator") == "bullet" else "  |  "
    for index, part in enumerate(parts[1:]):
        sep_run = paragraph.add_run(separator)
        sep_run.font.size = Pt(template_style["docx_body_size"])
        run = paragraph.add_run(part)
        run.font.size = Pt(template_style["docx_body_size"])
        if index == 0:
            run.italic = True

    if date_part:
        date_run = paragraph.add_run("\t" + date_part)
        date_run.font.size = Pt(template_style["docx_body_size"])


def _add_docx_body_paragraph(document: Document, text: str, template_style: dict) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(template_style["docx_item_space_after"])
    if ":" in text:
        label, rest = text.split(":", 1)
        if 1 <= len(label.split()) <= 4 and len(label) <= 32:
            label_run = paragraph.add_run(f"{label}:")
            label_run.bold = True
            label_run.font.size = Pt(template_style["docx_body_size"])
            rest_run = paragraph.add_run(rest)
            rest_run.font.size = Pt(template_style["docx_body_size"])
            return
    run = paragraph.add_run(text)
    run.font.size = Pt(template_style["docx_body_size"])


def _add_docx_bullet(document: Document, text: str, template_style: dict) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(template_style["docx_bullet_space_after"])
    paragraph.paragraph_format.left_indent = Inches(float(template_style.get("docx_bullet_left_indent", 0.24)))
    paragraph.paragraph_format.first_line_indent = Inches(float(template_style.get("docx_bullet_first_line_indent", -0.13)))
    run = paragraph.add_run(text)
    run.font.size = Pt(template_style["docx_body_size"])


def _add_docx_horizontal_line(document: Document, template_style: dict) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(template_style["docx_separator_space_after"])
    paragraph.paragraph_format.space_before = Pt(0)
    _add_docx_bottom_border(paragraph, template_style.get("separator_color", "B7B7B7"), "4", "1")


def _add_docx_bottom_border(paragraph, color: str, size: str, space: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), str(color))
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _docx_header_alignment(template_style: dict):
    if template_style.get("header_alignment") == "center":
        return WD_ALIGN_PARAGRAPH.CENTER
    return WD_ALIGN_PARAGRAPH.LEFT


def _pdf_horizontal_line(template_style: dict) -> HRFlowable:
    return HRFlowable(
        width="100%",
        thickness=0.45,
        lineCap="butt",
        color=colors.HexColor(f"#{template_style.get('separator_color', 'B7B7B7')}"),
        spaceBefore=0,
        spaceAfter=template_style["pdf_separator_space_after"],
    )


def _register_pdf_fonts() -> None:
    global PDF_FONT_NAME, PDF_BOLD_FONT_NAME

    if PDF_FONT_NAME != "Helvetica":
        return

    regular_paths = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    bold_paths = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
        "/Library/Fonts/Arial Bold.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    ]

    regular_path = _first_existing_path(regular_paths)
    bold_path = _first_existing_path(bold_paths)

    if regular_path:
        pdfmetrics.registerFont(TTFont("ATSCVFont", regular_path))
        PDF_FONT_NAME = "ATSCVFont"

    if bold_path:
        pdfmetrics.registerFont(TTFont("ATSCVFont-Bold", bold_path))
        PDF_BOLD_FONT_NAME = "ATSCVFont-Bold"
    elif regular_path:
        PDF_BOLD_FONT_NAME = "ATSCVFont"


def _pdf_styles(template_style: dict) -> dict:
    styles = getSampleStyleSheet()
    header_alignment = _pdf_header_alignment(template_style)
    return {
        "Name": ParagraphStyle(
            "ATSName",
            parent=styles["Normal"],
            fontName=PDF_BOLD_FONT_NAME,
            fontSize=template_style["pdf_name_size"],
            leading=template_style["pdf_name_size"] + 3,
            alignment=header_alignment,
            spaceAfter=3,
            textColor=colors.HexColor(f"#{template_style.get('accent_color', '111111')}"),
        ),
        "TargetTitle": ParagraphStyle(
            "ATSTargetTitle",
            parent=styles["Normal"],
            fontName=PDF_BOLD_FONT_NAME,
            fontSize=template_style["pdf_title_size"],
            leading=template_style["pdf_title_size"] + 2.5,
            alignment=header_alignment,
            spaceAfter=3,
            textColor=colors.HexColor(f"#{template_style.get('title_color', '111111')}"),
        ),
        "Contact": ParagraphStyle(
            "ATSContact",
            parent=styles["Normal"],
            fontName=PDF_FONT_NAME,
            fontSize=template_style["pdf_contact_size"],
            leading=template_style["pdf_contact_size"] + 2,
            alignment=header_alignment,
            spaceAfter=2.4,
            splitLongWords=0,
        ),
        "PhotoName": ParagraphStyle(
            "ATSPhotoName",
            parent=styles["Normal"],
            fontName=PDF_BOLD_FONT_NAME,
            fontSize=template_style["pdf_name_size"],
            leading=template_style["pdf_name_size"] + 3,
            alignment=TA_LEFT,
            spaceAfter=3,
            textColor=colors.HexColor(f"#{template_style.get('accent_color', '111111')}"),
        ),
        "PhotoTargetTitle": ParagraphStyle(
            "ATSPhotoTargetTitle",
            parent=styles["Normal"],
            fontName=PDF_BOLD_FONT_NAME,
            fontSize=template_style["pdf_title_size"],
            leading=template_style["pdf_title_size"] + 2.5,
            alignment=TA_LEFT,
            spaceAfter=3,
            textColor=colors.HexColor(f"#{template_style.get('title_color', '111111')}"),
        ),
        "PhotoContact": ParagraphStyle(
            "ATSPhotoContact",
            parent=styles["Normal"],
            fontName=PDF_FONT_NAME,
            fontSize=template_style["pdf_contact_size"],
            leading=template_style["pdf_contact_size"] + 2,
            alignment=TA_LEFT,
            spaceAfter=2,
            splitLongWords=0,
        ),
        "SectionHeading": ParagraphStyle(
            "ATSSectionHeading",
            parent=styles["Normal"],
            fontName=PDF_BOLD_FONT_NAME,
            fontSize=template_style["pdf_heading_size"],
            leading=template_style["pdf_heading_size"] + 2.5,
            alignment=TA_LEFT,
            spaceBefore=template_style["pdf_heading_space_before"],
            spaceAfter=template_style["pdf_heading_space_after"],
            textColor=colors.HexColor(f"#{template_style.get('heading_color', '111111')}"),
        ),
        "Separator": ParagraphStyle(
            "ATSSeparator",
            parent=styles["Normal"],
            fontName=PDF_FONT_NAME,
            fontSize=template_style["pdf_separator_size"],
            leading=template_style["pdf_separator_size"] + 1,
            spaceAfter=template_style["pdf_separator_space_after"],
        ),
        "ItemHeading": ParagraphStyle(
            "ATSItemHeading",
            parent=styles["Normal"],
            fontName=PDF_BOLD_FONT_NAME,
            fontSize=template_style["pdf_item_heading_size"],
            leading=template_style["pdf_item_heading_size"] + 2.5,
            spaceBefore=template_style.get("pdf_item_space_before", 2),
            spaceAfter=template_style.get("pdf_item_space_after", 1),
            textColor=colors.HexColor(f"#{template_style.get('item_heading_color', '111111')}"),
        ),
        "Body": ParagraphStyle(
            "ATSBody",
            parent=styles["Normal"],
            fontName=PDF_FONT_NAME,
            fontSize=template_style["pdf_body_size"],
            leading=template_style["pdf_body_size"] + 2.5,
            spaceAfter=template_style["pdf_body_space_after"],
        ),
        "Bullet": ParagraphStyle(
            "ATSBullet",
            parent=styles["Normal"],
            fontName=PDF_FONT_NAME,
            fontSize=template_style["pdf_body_size"],
            leading=template_style["pdf_body_size"] + 2.3,
            leftIndent=template_style.get("pdf_bullet_left_indent", 12),
            firstLineIndent=template_style.get("pdf_bullet_first_line_indent", -8),
            spaceAfter=template_style["pdf_bullet_space_after"],
        ),
    }


def _pdf_header(
    contact: dict,
    styles: dict,
    template_style: dict,
    photo_bytes: bytes | None = None,
    photo_filename: str = "",
) -> list:
    flowables = []
    if template_style.get("supports_photo") and photo_bytes:
        photo_header = _pdf_photo_header(contact, styles, template_style, photo_bytes, photo_filename)
        if photo_header:
            flowables.extend(photo_header)
        else:
            flowables.extend(_pdf_text_header(contact, styles))
    else:
        flowables.extend(_pdf_text_header(contact, styles))

    if flowables and template_style.get("header_separator"):
        flowables.append(_pdf_horizontal_line(template_style))
        flowables.append(Spacer(1, template_style.get("pdf_header_rule_after_spacing", template_style["pdf_contact_after_spacing"])))
    elif flowables:
        flowables.append(Spacer(1, template_style["pdf_contact_after_spacing"]))
    return flowables


def _pdf_text_header(contact: dict, styles: dict) -> list:
    flowables = []
    if contact["full_name"]:
        flowables.append(Paragraph(escape(contact["full_name"]), styles["Name"]))
    if contact["target_title"]:
        flowables.append(Paragraph(escape(contact["target_title"]), styles["TargetTitle"]))
    for contact_line in contact["contact_lines"]:
        flowables.append(Paragraph(escape(contact_line), styles["Contact"]))
    return flowables


def _pdf_photo_header(
    contact: dict,
    styles: dict,
    template_style: dict,
    photo_bytes: bytes,
    photo_filename: str,
) -> list:
    try:
        photo_bytes = prepare_cv_photo_for_export(photo_bytes)
        photo_size = float(template_style.get("pdf_photo_size", 0.82)) * inch
        photo = RLImage(BytesIO(photo_bytes), width=photo_size, height=photo_size)
    except Exception:
        return []

    text_flowables = []
    if contact["full_name"]:
        text_flowables.append(Paragraph(escape(contact["full_name"]), styles["PhotoName"]))
    if contact["target_title"]:
        text_flowables.append(Paragraph(escape(contact["target_title"]), styles["PhotoTargetTitle"]))
    for contact_line in contact["contact_lines"]:
        text_flowables.append(Paragraph(escape(contact_line), styles["PhotoContact"]))

    table = Table(
        [[photo, text_flowables]],
        colWidths=[
            float(template_style.get("pdf_photo_cell_width", 0.98)) * inch,
            float(template_style.get("pdf_photo_text_width", 6.05)) * inch,
        ],
        hAlign="LEFT",
    )
    table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (0, 0), 8),
        ("RIGHTPADDING", (1, 0), (1, 0), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    return [table]


def _pdf_item_heading_markup(text: str, template_style: dict) -> str:
    parts = [part.strip() for part in str(text or "").split(" | ") if part.strip()]
    if not parts:
        return ""
    if template_style.get("item_separator") != "bullet":
        return escape(text)

    date_part = ""
    if len(parts) >= 2 and _looks_like_date_range(parts[-1]):
        date_part = parts[-1]
        parts = parts[:-1]

    separator = " &bull; "
    markup_parts = [f"<b>{escape(parts[0])}</b>"]
    for index, part in enumerate(parts[1:]):
        escaped_part = escape(part)
        markup_parts.append(f"<i>{escaped_part}</i>" if index == 0 else escaped_part)
    if date_part:
        markup_parts.append(escape(date_part))
    return separator.join(markup_parts)


def _pdf_header_alignment(template_style: dict) -> int:
    if template_style.get("header_alignment") == "center":
        return TA_CENTER
    return TA_LEFT


def _language_key(language: str) -> str:
    return "Turkish" if str(language).lower() == "turkish" else "English"


def _section_title(section_key: str, language: str) -> str:
    lang = _language_key(language)
    return SECTION_TITLES[lang].get(section_key, section_key.replace("_", " ").title())


def _heading_text(text: str, language: str) -> str:
    if _language_key(language) == "Turkish":
        return text.translate(str.maketrans({"i": "İ", "ı": "I"})).upper()
    return text.upper()


def _skill_group_title(group: str, language: str) -> str:
    titles = {
        "English": {
            "technical_skills": "Technical Skills",
            "core_skills": "Core Skills",
            "tools": "Tools",
            "databases": "Databases",
            "cloud": "Cloud",
            "soft_skills": "Soft Skills",
        },
        "Turkish": {
            "technical_skills": "Teknik Yetenekler",
            "core_skills": "Temel Yetkinlikler",
            "tools": "Araçlar",
            "databases": "Veritabanları",
            "cloud": "Bulut",
            "soft_skills": "Sosyal Beceriler",
        },
    }
    return titles[_language_key(language)].get(group, group.replace("_", " ").title())


def _clean_text(value) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _clean_list(values) -> list[str]:
    if not isinstance(values, list):
        return []
    return [_clean_text(value) for value in values if _clean_text(value)]


def _modern_clean_core_skill_values(skills: dict, metadata: dict | None = None) -> list[str]:
    core_values = _clean_list(skills.get("core_skills", []))
    all_values = core_values[:]
    for group in ["technical_skills", "tools", "databases", "cloud", "soft_skills"]:
        all_values.extend(_clean_list(skills.get(group, [])))

    ranked_values = rank_skills_for_job({"core": _dedupe_preserve_order(all_values)}, "", metadata or {}).get("core", [])
    return _dedupe_preserve_order(ranked_values + core_values)[:16]


def _join_non_empty(values, separator: str = " | ") -> str:
    return separator.join(_clean_text(value) for value in values if _clean_text(value))


def _date_range(start_date, end_date) -> str:
    return _join_non_empty([start_date, end_date], separator=" - ")


def _looks_like_date_range(value: str) -> bool:
    text = _clean_text(value).lower()
    if not text:
        return False
    return bool(re.search(r"\b(19|20)\d{2}\b|present|current|devam|günümüz|ocak|şubat|mart|nisan|mayıs|haziran|temmuz|ağustos|eylül|ekim|kasım|aralık|jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec", text))


def _project_link_label(link: str, contact_links: set[str], language: str) -> str:
    cleaned_link = _clean_text(link)
    if not cleaned_link:
        return ""

    normalized_link = _normalize_url_for_compare(cleaned_link)
    if normalized_link in contact_links or _is_github_profile_link(cleaned_link):
        return ""

    if "github.com/" in normalized_link:
        label = "GitHub" if _language_key(language) == "English" else "GitHub"
    else:
        label = "Project Link" if _language_key(language) == "English" else "Proje Linki"
    return f"{label}: {cleaned_link}"


def _is_github_profile_link(link: str) -> bool:
    normalized = _normalize_url_for_compare(link)
    if "github.com/" not in normalized:
        return False
    path = normalized.split("github.com/", 1)[1].strip("/")
    return bool(path) and "/" not in path


def _normalize_url_for_compare(value) -> str:
    text = _clean_text(value).lower().rstrip("/")
    for prefix in ["https://www.", "http://www.", "https://", "http://", "www."]:
        if text.startswith(prefix):
            return text[len(prefix):]
    return text


def _trim_text(value, max_chars: int) -> str:
    text = _clean_text(value)
    if len(text) <= max_chars:
        return text

    trimmed = text[:max_chars].rsplit(" ", 1)[0].strip()
    return f"{trimmed}."


def _prioritize_strings(values: list[str], keywords: list[str], limit: int) -> list[str]:
    if not values:
        return []
    if limit <= 0:
        return []

    normalized_keywords = _dedupe_preserve_order([
        keyword.lower().strip()
        for keyword in keywords
        if _clean_text(keyword)
    ])

    def score(value: str) -> int:
        normalized_value = value.lower()
        value_tokens = set(re.findall(r"[\w#+.]+", normalized_value))
        total = 0
        for keyword in normalized_keywords:
            if keyword in normalized_value:
                total += 4
            keyword_tokens = set(re.findall(r"[\w#+.]+", keyword))
            if keyword_tokens and value_tokens:
                total += len(keyword_tokens & value_tokens)
        return total

    indexed_values = list(enumerate(values))
    sorted_values = sorted(indexed_values, key=lambda item: (-score(item[1]), item[0]))
    selected = sorted_values[:limit]
    return [value for _, value in sorted(selected, key=lambda item: item[0])]


def _dedupe_preserve_order(values: list[str]) -> list[str]:
    seen = set()
    result = []
    for value in values:
        cleaned = _clean_text(value)
        key = cleaned.lower()
        if cleaned and key not in seen:
            seen.add(key)
            result.append(cleaned)
    return result


def _cv_for_export_style(ats_cv: dict, template: dict, language: str, export_style: str) -> dict:
    from services.ats_cv_postprocessing import clean_structured_cv_before_export
    ats_cv = clean_structured_cv_before_export(ats_cv)
    original_cv = deepcopy(ats_cv)
    if export_style in {"compact", "balanced_one_page"}:
        render_cv = balance_one_page_content(ats_cv, template, language)
    else:
        render_cv = deepcopy(ats_cv)

    return _restore_preserved_export_fields(original_cv, render_cv)


def _restore_preserved_export_fields(source_cv: dict, render_cv: dict) -> dict:
    """Ensure export styling and compaction never rewrite protected identity fields."""
    result = deepcopy(render_cv)
    if not isinstance(source_cv, dict):
        return result

    source_contact = source_cv.get("contact")
    result_contact = result.get("contact")
    if isinstance(source_contact, dict) and isinstance(result_contact, dict):
        for field in PRESERVED_CONTACT_FIELDS:
            if _clean_text(source_contact.get(field)):
                result_contact[field] = source_contact.get(field)

    for section_key, fields in PRESERVED_RECORD_FIELDS.items():
        source_records = source_cv.get(section_key)
        result_records = result.get(section_key)
        if not isinstance(source_records, list) or not isinstance(result_records, list):
            continue

        for index, source_record in enumerate(source_records):
            if index >= len(result_records):
                break
            result_record = result_records[index]
            if not isinstance(source_record, dict) or not isinstance(result_record, dict):
                continue
            for field in fields:
                if _clean_text(source_record.get(field)):
                    result_record[field] = source_record.get(field)

    return result


def _effective_export_style(export_style: str, one_page: bool) -> str:
    normalized = str(export_style or "").strip().lower().replace(" ", "_").replace("-", "_")
    if normalized in {"", "standard", "normal"}:
        normalized = "balanced_one_page" if one_page else "standard"
    if normalized in {"one_page", "balanced", "balanced_one_page"}:
        return "balanced_one_page"
    if normalized == "compact":
        return "compact"
    return "standard"


def _density_limits(density: str) -> dict:
    if density == "short":
        return {
            "summary_chars": 760,
            "experience_bullets": 4,
            "project_bullets": 3,
            "project_description_chars": 360,
            "project_technologies": 12,
            "education_details": 3,
            "certifications": 6,
            "skills": {
                "technical_skills": 16,
                "core_skills": 12,
                "tools": 8,
                "databases": 8,
                "cloud": 6,
                "soft_skills": 6,
            },
        }
    if density == "medium":
        return {
            "summary_chars": 620,
            "experience_bullets": 3,
            "project_bullets": 2,
            "project_description_chars": 300,
            "project_technologies": 10,
            "education_details": 2,
            "certifications": 4,
            "skills": {
                "technical_skills": 13,
                "core_skills": 10,
                "tools": 7,
                "databases": 6,
                "cloud": 5,
                "soft_skills": 5,
            },
        }
    return {
        "summary_chars": 460,
        "experience_bullets": 2,
        "project_bullets": 2,
        "project_description_chars": 220,
        "project_technologies": 8,
        "education_details": 1,
        "certifications": 3,
        "skills": {
            "technical_skills": 10,
            "core_skills": 8,
            "tools": 5,
            "databases": 5,
            "cloud": 4,
            "soft_skills": 4,
        },
    }


def _first_existing_path(paths: list[str]) -> str:
    for path in paths:
        if Path(path).exists():
            return path
    return ""


def _docx_rgb(value: str) -> RGBColor:
    text = str(value or "111111").strip().lstrip("#")
    if len(text) != 6 or any(character not in "0123456789abcdefABCDEF" for character in text):
        text = "111111"
    return RGBColor(int(text[0:2], 16), int(text[2:4], 16), int(text[4:6], 16))


def _template_style(template: dict, export_style: str = "standard", density: str = "medium") -> dict:
    template_id = template.get("id", "classic_ats")
    styles = {
        "classic_ats": {
            "docx_margin": 0.58,
            "docx_name_size": 18,
            "docx_title_size": 11,
            "docx_contact_size": 9,
            "docx_body_size": 10.4,
            "docx_section_space_before": 8,
            "docx_item_space_after": 1.8,
            "docx_bullet_space_after": 1,
            "pdf_margin": 0.58,
            "pdf_name_size": 17,
            "pdf_title_size": 11,
            "pdf_contact_size": 8.7,
            "pdf_heading_size": 10.6,
            "pdf_item_heading_size": 9.5,
            "pdf_body_size": 9.25,
            "pdf_body_space_after": 1.8,
            "pdf_bullet_space_after": 1,
            "pdf_contact_after_spacing": 0.14 * inch,
            "pdf_section_spacing": 0,
            "pdf_section_after_spacing": 0.07 * inch,
            "heading_separator": False,
        },
        "modern_clean": {
            "docx_margin": 0.54,
            "docx_name_size": 18,
            "docx_title_size": 11.5,
            "docx_contact_size": 8.8,
            "docx_body_size": 9.6,
            "docx_section_space_before": 7,
            "docx_item_space_after": 1.1,
            "docx_bullet_space_after": 0.55,
            "docx_heading_size": 10.3,
            "docx_heading_space_after": 0.4,
            "docx_separator_size": 7.5,
            "docx_separator_space_after": 0.5,
            "pdf_margin": 0.52,
            "pdf_name_size": 17.8,
            "pdf_title_size": 11.5,
            "pdf_contact_size": 8.8,
            "pdf_heading_size": 10.2,
            "pdf_item_heading_size": 9.6,
            "pdf_body_size": 9.0,
            "pdf_body_space_after": 0.8,
            "pdf_bullet_space_after": 0.35,
            "pdf_heading_space_before": 3.2,
            "pdf_heading_space_after": 0.6,
            "pdf_separator_size": 5.6,
            "pdf_separator_space_after": 0.45,
            "pdf_contact_after_spacing": 0.09 * inch,
            "pdf_section_spacing": 0,
            "pdf_section_after_spacing": 0.045 * inch,
            "heading_separator": True,
            "separator_text": "-" * 54,
            "separator_as_rule": True,
        },
        "technical_developer": {
            "docx_margin": 0.56,
            "docx_name_size": 17,
            "docx_title_size": 10.5,
            "docx_contact_size": 8.5,
            "docx_body_size": 10,
            "docx_section_space_before": 7,
            "docx_item_space_after": 1.2,
            "docx_bullet_space_after": 0.7,
            "pdf_margin": 0.54,
            "pdf_name_size": 17,
            "pdf_title_size": 10.8,
            "pdf_contact_size": 8.2,
            "pdf_heading_size": 10.4,
            "pdf_item_heading_size": 9.5,
            "pdf_body_size": 9.05,
            "pdf_body_space_after": 1.2,
            "pdf_bullet_space_after": 0.6,
            "pdf_contact_after_spacing": 0.10 * inch,
            "pdf_section_spacing": 0,
            "pdf_section_after_spacing": 0.06 * inch,
            "heading_separator": False,
        },
        "junior_internship": {
            "docx_margin": 0.58,
            "docx_name_size": 17,
            "docx_title_size": 10.8,
            "docx_contact_size": 8.7,
            "docx_body_size": 10.2,
            "docx_section_space_before": 9,
            "docx_item_space_after": 1.6,
            "docx_bullet_space_after": 0.9,
            "pdf_margin": 0.58,
            "pdf_name_size": 17,
            "pdf_title_size": 10.8,
            "pdf_contact_size": 8.5,
            "pdf_heading_size": 10.5,
            "pdf_item_heading_size": 9.3,
            "pdf_body_size": 9,
            "pdf_body_space_after": 1.4,
            "pdf_bullet_space_after": 0.8,
            "pdf_contact_after_spacing": 0.12 * inch,
            "pdf_section_spacing": 0,
            "pdf_section_after_spacing": 0.09 * inch,
            "heading_separator": False,
        },
        "modern_professional": {
            "docx_margin": 0.56,
            "docx_name_size": 24.5,
            "docx_title_size": 11.8,
            "docx_contact_size": 8.1,
            "docx_body_size": 9.55,
            "docx_section_space_before": 8.4,
            "docx_item_space_before": 2.8,
            "docx_item_space_after": 0.85,
            "docx_bullet_space_after": 0.35,
            "docx_heading_size": 10.8,
            "docx_heading_space_after": 0.8,
            "docx_separator_size": 7,
            "docx_separator_space_after": 0.7,
            "docx_header_rule_space_after": 7,
            "docx_bullet_left_indent": 0.22,
            "docx_bullet_first_line_indent": -0.13,
            "docx_date_tab_inch": 6.52,
            "pdf_margin": 0.54,
            "pdf_name_size": 22.5,
            "pdf_title_size": 11.6,
            "pdf_contact_size": 8.0,
            "pdf_heading_size": 10.4,
            "pdf_item_heading_size": 9.25,
            "pdf_item_space_before": 2.4,
            "pdf_item_space_after": 0.8,
            "pdf_body_size": 8.8,
            "pdf_body_space_after": 0.55,
            "pdf_bullet_space_after": 0.18,
            "pdf_bullet_left_indent": 11,
            "pdf_bullet_first_line_indent": -7,
            "pdf_heading_space_before": 3.2,
            "pdf_heading_space_after": 0.7,
            "pdf_separator_size": 5.8,
            "pdf_separator_space_after": 0.45,
            "pdf_contact_after_spacing": 0.08 * inch,
            "pdf_header_rule_after_spacing": 0.055 * inch,
            "pdf_section_spacing": 0,
            "pdf_section_after_spacing": 0.04 * inch,
            "heading_separator": True,
            "header_separator": True,
            "header_alignment": "center",
            "item_separator": "bullet",
            "separator_text": "-" * 54,
            "separator_as_rule": True,
            "separator_color": "AEBAC6",
            "accent_color": "274B63",
            "title_color": "333333",
            "heading_color": "274B63",
            "item_heading_color": "111827",
            "contact_max_chars": 104,
            "shorten_contact_links": True,
        },
        "compact_technical": {
            "docx_margin": 0.48,
            "docx_name_size": 19,
            "docx_title_size": 10.2,
            "docx_contact_size": 7.55,
            "docx_body_size": 8.65,
            "docx_section_space_before": 5.2,
            "docx_item_space_before": 1.4,
            "docx_item_space_after": 0.2,
            "docx_bullet_space_after": 0,
            "docx_heading_size": 9.8,
            "docx_heading_space_after": 0.25,
            "docx_separator_size": 6.5,
            "docx_separator_space_after": 0.25,
            "docx_header_rule_space_after": 4.5,
            "docx_bullet_left_indent": 0.18,
            "docx_bullet_first_line_indent": -0.11,
            "docx_date_tab_inch": 6.82,
            "pdf_margin": 0.48,
            "pdf_name_size": 18,
            "pdf_title_size": 10.2,
            "pdf_contact_size": 7.45,
            "pdf_heading_size": 9.6,
            "pdf_item_heading_size": 8.65,
            "pdf_item_space_before": 1.4,
            "pdf_item_space_after": 0.15,
            "pdf_body_size": 8.15,
            "pdf_body_space_after": 0.12,
            "pdf_bullet_space_after": 0,
            "pdf_bullet_left_indent": 9,
            "pdf_bullet_first_line_indent": -6,
            "pdf_heading_space_before": 2.2,
            "pdf_heading_space_after": 0.2,
            "pdf_separator_size": 5.4,
            "pdf_separator_space_after": 0.1,
            "pdf_contact_after_spacing": 0.045 * inch,
            "pdf_header_rule_after_spacing": 0.035 * inch,
            "pdf_section_spacing": 0,
            "pdf_section_after_spacing": 0.02 * inch,
            "heading_separator": True,
            "header_separator": True,
            "header_alignment": "left",
            "separator_text": "-" * 58,
            "separator_as_rule": True,
            "separator_color": "222222",
            "accent_color": "111111",
            "title_color": "111111",
            "heading_color": "111111",
            "item_heading_color": "111111",
            "contact_max_chars": 112,
            "shorten_contact_links": True,
        },
        "visual_photo_optional": {
            "docx_margin": 0.56,
            "docx_name_size": 23.5,
            "docx_title_size": 11.4,
            "docx_contact_size": 8.0,
            "docx_body_size": 9.3,
            "docx_section_space_before": 7.8,
            "docx_item_space_before": 2.3,
            "docx_item_space_after": 0.7,
            "docx_bullet_space_after": 0.28,
            "docx_heading_size": 10.55,
            "docx_heading_space_after": 0.6,
            "docx_separator_size": 6.8,
            "docx_separator_space_after": 0.6,
            "docx_header_rule_space_after": 6,
            "docx_bullet_left_indent": 0.21,
            "docx_bullet_first_line_indent": -0.12,
            "docx_date_tab_inch": 6.48,
            "docx_photo_width": 1.04,
            "docx_photo_cell_width": 1.22,
            "pdf_margin": 0.54,
            "pdf_name_size": 21.0,
            "pdf_title_size": 11.2,
            "pdf_contact_size": 7.95,
            "pdf_heading_size": 10.25,
            "pdf_item_heading_size": 9.2,
            "pdf_item_space_before": 2.1,
            "pdf_item_space_after": 0.6,
            "pdf_body_size": 8.65,
            "pdf_body_space_after": 0.45,
            "pdf_bullet_space_after": 0.15,
            "pdf_bullet_left_indent": 10,
            "pdf_bullet_first_line_indent": -7,
            "pdf_heading_space_before": 3.0,
            "pdf_heading_space_after": 0.55,
            "pdf_separator_size": 5.6,
            "pdf_separator_space_after": 0.4,
            "pdf_contact_after_spacing": 0.075 * inch,
            "pdf_header_rule_after_spacing": 0.05 * inch,
            "pdf_photo_size": 0.92,
            "pdf_photo_cell_width": 1.06,
            "pdf_photo_text_width": 6.05,
            "pdf_section_spacing": 0,
            "pdf_section_after_spacing": 0.035 * inch,
            "heading_separator": True,
            "header_separator": True,
            "header_alignment": "left",
            "item_separator": "bullet",
            "separator_text": "-" * 54,
            "separator_as_rule": True,
            "separator_color": "8A8F98",
            "accent_color": "1F2933",
            "title_color": "333333",
            "heading_color": "1F2933",
            "item_heading_color": "111827",
            "supports_photo": True,
            "contact_max_chars": 98,
            "shorten_contact_links": True,
            "contact_separator": "  •  ",
            "italic_title": True,
        },
    }
    style = deepcopy(styles.get(template_id, styles["classic_ats"]))
    style["template_id"] = template_id
    style.setdefault("docx_contact_size", 9)
    style.setdefault("docx_heading_size", 11)
    style.setdefault("docx_heading_space_after", 3)
    style.setdefault("docx_separator_size", 8)
    style.setdefault("docx_separator_space_after", 3)
    style.setdefault("docx_header_rule_space_after", 6)
    style.setdefault("docx_item_space_before", 1.5)
    style.setdefault("docx_bullet_left_indent", 0.24)
    style.setdefault("docx_bullet_first_line_indent", -0.13)
    style.setdefault("docx_date_tab_inch", 6.45)
    style.setdefault("pdf_heading_space_before", 6)
    style.setdefault("pdf_heading_space_after", 2.5)
    style.setdefault("pdf_separator_size", 7)
    style.setdefault("pdf_separator_space_after", 3)
    style.setdefault("pdf_header_rule_after_spacing", style.get("pdf_contact_after_spacing", 0.08 * inch))
    style.setdefault("pdf_item_space_before", 2)
    style.setdefault("pdf_item_space_after", 1)
    style.setdefault("pdf_bullet_left_indent", 12)
    style.setdefault("pdf_bullet_first_line_indent", -8)
    style.setdefault("separator_text", "-" * 48)
    style.setdefault("separator_as_rule", False)
    style.setdefault("header_separator", False)
    style.setdefault("header_alignment", "center")
    style.setdefault("item_separator", "pipe")
    style.setdefault("item_heading_color", "111111")
    style.setdefault("contact_max_chars", 104)
    style.setdefault("shorten_contact_links", False)
    style.setdefault("contact_separator", " | ")
    style.setdefault("italic_title", False)
    style.setdefault("supports_photo", False)
    if export_style == "compact":
        _apply_compact_style(style)
    elif export_style == "balanced_one_page":
        _apply_balanced_one_page_style(style, density)
    return style


def _apply_compact_style(style: dict) -> None:
    style["docx_margin"] = max(0.45, style["docx_margin"] - 0.14)
    style["docx_name_size"] = max(15, style["docx_name_size"] - 1.4)
    style["docx_title_size"] = max(9.5, style["docx_title_size"] - 0.9)
    style["docx_body_size"] = max(9, style["docx_body_size"] - 0.8)
    style["docx_heading_size"] = max(10, style["docx_heading_size"] - 0.4)
    style["docx_heading_space_after"] = max(1, style["docx_heading_space_after"] - 1.2)
    style["docx_separator_space_after"] = max(0.5, style["docx_separator_space_after"] - 1.5)
    style["docx_section_space_before"] = max(4, style["docx_section_space_before"] - 4)
    style["docx_item_space_after"] = max(0, style["docx_item_space_after"] - 1.2)
    style["docx_bullet_space_after"] = max(0, style["docx_bullet_space_after"] - 0.8)
    style["pdf_margin"] = max(0.45, style["pdf_margin"] - 0.12)
    style["pdf_name_size"] = max(15, style["pdf_name_size"] - 1.2)
    style["pdf_title_size"] = max(10, style["pdf_title_size"] - 0.8)
    style["pdf_contact_size"] = max(8, style["pdf_contact_size"] - 0.5)
    style["pdf_heading_size"] = max(9.5, style["pdf_heading_size"] - 0.8)
    style["pdf_heading_space_before"] = max(2, style["pdf_heading_space_before"] - 3)
    style["pdf_heading_space_after"] = max(0.8, style["pdf_heading_space_after"] - 1.2)
    style["pdf_separator_size"] = max(5.8, style["pdf_separator_size"] - 0.6)
    style["pdf_separator_space_after"] = max(0, style["pdf_separator_space_after"] - 2)
    style["pdf_item_heading_size"] = max(8.7, style["pdf_item_heading_size"] - 0.5)
    style["pdf_body_size"] = max(8.5, style["pdf_body_size"] - 0.6)
    style["pdf_body_space_after"] = max(0, style["pdf_body_space_after"] - 1)
    style["pdf_bullet_space_after"] = max(0, style["pdf_bullet_space_after"] - 0.7)
    style["pdf_contact_after_spacing"] = max(0.06 * inch, style["pdf_contact_after_spacing"] - 0.05 * inch)
    style["pdf_section_spacing"] = max(0, style["pdf_section_spacing"] - 0.03 * inch)
    style["pdf_section_after_spacing"] = max(0.03 * inch, style["pdf_section_after_spacing"] - 0.05 * inch)


def _apply_balanced_one_page_style(style: dict, density: str) -> None:
    if density == "short":
        style["docx_margin"] = max(0.50, style["docx_margin"] - 0.06)
        style["pdf_margin"] = max(0.50, style["pdf_margin"] - 0.06)
        style["docx_section_space_before"] += 1
        style["pdf_section_after_spacing"] += 0.02 * inch
        style["pdf_body_size"] = min(9.5, style["pdf_body_size"] + 0.15)
        style["pdf_bullet_space_after"] += 0.2
        _apply_modern_clean_one_page_micro_adjustments(style, density)
        return

    if density == "medium":
        style["docx_margin"] = max(0.50, style["docx_margin"] - 0.08)
        style["docx_name_size"] = max(15.5, style["docx_name_size"] - 0.5)
        style["docx_body_size"] = max(9.4, style["docx_body_size"] - 0.35)
        style["docx_section_space_before"] = max(5, style["docx_section_space_before"] - 2)
        style["docx_heading_space_after"] = max(1.5, style["docx_heading_space_after"] - 0.5)
        style["pdf_margin"] = max(0.50, style["pdf_margin"] - 0.08)
        style["pdf_name_size"] = max(15.5, style["pdf_name_size"] - 0.6)
        style["pdf_title_size"] = max(10.3, style["pdf_title_size"] - 0.4)
        style["pdf_heading_size"] = max(9.8, style["pdf_heading_size"] - 0.4)
        style["pdf_heading_space_before"] = max(3, style["pdf_heading_space_before"] - 1.2)
        style["pdf_heading_space_after"] = max(1.2, style["pdf_heading_space_after"] - 0.6)
        style["pdf_body_size"] = max(8.7, style["pdf_body_size"] - 0.3)
        style["pdf_body_space_after"] = max(0.4, style["pdf_body_space_after"] - 0.4)
        style["pdf_bullet_space_after"] = max(0.2, style["pdf_bullet_space_after"] - 0.25)
        style["pdf_section_after_spacing"] = max(0.05 * inch, style["pdf_section_after_spacing"] - 0.03 * inch)
        _apply_modern_clean_one_page_micro_adjustments(style, density)
        return

    _apply_compact_style(style)
    _apply_modern_clean_one_page_micro_adjustments(style, density)


def _apply_modern_clean_one_page_micro_adjustments(style: dict, density: str) -> None:
    if style.get("template_id") != "modern_clean":
        return

    reduction = 0.16 if density == "short" else 0.32
    style["docx_margin"] = max(0.46, style["docx_margin"] - 0.03)
    style["docx_body_size"] = max(9.25, style["docx_body_size"] - reduction)
    style["docx_section_space_before"] = max(4.2, style["docx_section_space_before"] - 1.2)
    style["docx_heading_space_after"] = max(0.8, style["docx_heading_space_after"] - 0.8)
    style["docx_separator_space_after"] = max(0.3, style["docx_separator_space_after"] - 1.8)
    style["docx_item_space_after"] = max(0.3, style["docx_item_space_after"] - 0.6)
    style["docx_bullet_space_after"] = max(0.15, style["docx_bullet_space_after"] - 0.45)

    style["pdf_margin"] = max(0.44, style["pdf_margin"] - 0.045)
    style["pdf_body_size"] = max(8.45, style["pdf_body_size"] - reduction)
    style["pdf_heading_size"] = max(9.45, style["pdf_heading_size"] - 0.2)
    style["pdf_heading_space_before"] = max(1.5, style["pdf_heading_space_before"] - 1.6)
    style["pdf_heading_space_after"] = max(0.6, style["pdf_heading_space_after"] - 0.8)
    style["pdf_separator_size"] = max(5.6, style["pdf_separator_size"] - 0.8)
    style["pdf_separator_space_after"] = max(0, style["pdf_separator_space_after"] - 2.4)
    style["pdf_body_space_after"] = max(0, style["pdf_body_space_after"] - 0.7)
    style["pdf_bullet_space_after"] = max(0, style["pdf_bullet_space_after"] - 0.45)
    style["pdf_contact_after_spacing"] = max(0.045 * inch, style["pdf_contact_after_spacing"] - 0.035 * inch)
    style["pdf_section_spacing"] = max(0, style["pdf_section_spacing"] - 0.025 * inch)
    style["pdf_section_after_spacing"] = max(0.015 * inch, style["pdf_section_after_spacing"] - 0.045 * inch)
