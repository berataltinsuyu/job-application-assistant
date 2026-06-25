import json
import sys
import tempfile
import zipfile
from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from services.ats_cv_templates import get_ats_cv_templates  # noqa: E402
from services.ats_cv_schema import get_empty_ats_cv_schema  # noqa: E402
from services.docx_template_service import get_docx_template_catalog, render_cv_with_docx_template  # noqa: E402
from services.llm_service import build_ats_cv_generation_prompt  # noqa: E402
from routers.ats_cv_builder import _neutralize_unsupported_student_wording  # noqa: E402
from main import app  # noqa: E402


NEW_TEMPLATE_IDS = {"modern_professional", "compact_technical", "visual_photo_optional"}
OLD_TEMPLATE_IDS = {"ats_classic_docx", "ats_modern_docx"}


def assert_true(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def sample_cv() -> dict:
    cv = get_empty_ats_cv_schema()
    cv.update({
        "contact": {
            "full_name": "Taylor Applicant",
            "target_title": "Product-Oriented Backend Developer",
            "email": "taylor@example.com",
            "phone": "+1 555 010 2222",
            "location": "Remote",
            "linkedin": "https://linkedin.com/in/taylorapplicant",
            "github": "https://github.com/taylorapplicant",
            "portfolio": "https://taylor.example.com",
        },
        "professional_summary": (
            "Backend developer with project-based API, SQL, automation, documentation, "
            "and product collaboration experience."
        ),
        "technical_summary": "Python, FastAPI, SQL, REST APIs, Docker, testing, documentation.",
        "skills": {
            "Backend": ["Python", "FastAPI", "REST APIs"],
            "Database": ["PostgreSQL", "SQLite"],
            "Tools": ["Git", "Docker", "Postman"],
            "Core Skills": ["Documentation", "Stakeholder communication"],
        },
        "experience": [
            {
                "title": "Backend Developer Intern",
                "company": "Example Software Studio",
                "location": "Remote",
                "start_date": "2025",
                "end_date": "2026",
                "bullets": [
                    "Built API endpoints for internal workflow tools.",
                    "Documented test cases, integration behavior, and release notes.",
                ],
            }
        ],
        "projects": [
            {
                "name": "Application Materials Portal",
                "technologies": ["Python", "FastAPI", "SQLite"],
                "description": "Built a small document generation and tracking prototype.",
                "bullets": ["Implemented validation, export, and preview behavior."],
            }
        ],
        "education": [
            {
                "school": "Example University",
                "degree": "BS Computer Engineering",
                "start_date": "2021",
                "end_date": "2025",
                "details": ["Relevant coursework in databases and software engineering."],
            }
        ],
        "certifications": [{"name": "Python Foundations", "issuer": "Example Academy", "date": "2025"}],
        "languages": [{"language": "English", "level": "Professional"}],
    })
    return cv


def sample_photo_bytes(width: int = 120, height: int = 120) -> bytes:
    image = Image.new("RGB", (width, height), color=(64, 92, 128))
    output = BytesIO()
    image.save(output, format="PNG")
    return output.getvalue()


def docx_text(path: Path) -> str:
    document = Document(str(path))
    lines = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                lines.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(lines)


def assert_social_contact_line(text: str, template_id: str) -> None:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    assert_true("linkedin.com/in/taylorapplicant" in text, f"{template_id} dropped or corrupted LinkedIn contact.")
    assert_true("github.com/taylorapplicant" in text, f"{template_id} dropped or corrupted GitHub contact.")
    assert_true(
        any("linkedin.com/in/taylorapplicant" in line and "github.com/taylorapplicant" in line for line in lines),
        f"{template_id} should keep LinkedIn and GitHub together in the contact block.",
    )


def docx_has_media(path: Path) -> bool:
    with zipfile.ZipFile(path) as archive:
        return any(name.startswith("word/media/") for name in archive.namelist())


def first_docx_media_size(path: Path) -> tuple[int, int]:
    with zipfile.ZipFile(path) as archive:
        media_names = [name for name in archive.namelist() if name.startswith("word/media/")]
        assert_true(bool(media_names), f"No media found in {path}")
        with Image.open(BytesIO(archive.read(media_names[0]))) as image:
            return image.size


def main() -> None:
    ats_ids = {template["id"] for template in get_ats_cv_templates()}
    catalog = get_docx_template_catalog()
    docx_ids = {template["template_id"] for template in catalog}
    catalog_by_id = {template["template_id"]: template for template in catalog}

    assert_true(NEW_TEMPLATE_IDS.issubset(ats_ids), "New visual ATS templates are missing from ATS catalog.")
    assert_true(NEW_TEMPLATE_IDS.issubset(docx_ids), "New visual DOCX templates are missing from DOCX catalog.")
    assert_true(OLD_TEMPLATE_IDS.issubset(docx_ids), "Old DOCX templates are missing from DOCX catalog.")
    assert_true(catalog_by_id["visual_photo_optional"].get("supports_photo") is True, "Photo template must advertise photo support.")
    assert_true(catalog_by_id["modern_professional"].get("supports_photo") is False, "Non-photo template should not advertise photo support.")

    cv = sample_cv()
    forbidden_reference_terms = ["Jacqueline Thompson", "Helin Kinay", "Borcelle Technologies", "XarrowAI"]

    with tempfile.TemporaryDirectory(dir=ROOT / "scratch") as tmp_dir:
        tmp_path = Path(tmp_dir)
        for template_id in sorted(OLD_TEMPLATE_IDS | NEW_TEMPLATE_IDS):
            output_path = tmp_path / f"{template_id}.docx"
            result = render_cv_with_docx_template(
                structured_cv=cv,
                template_id=template_id,
                output_path=str(output_path),
                metadata={"smoke_test": True},
            )
            assert_true(result["success"] is True, f"{template_id} render failed: {result}")
            assert_true(output_path.exists(), f"{template_id} output was not created.")
            assert_true(output_path.stat().st_size > 10_000, f"{template_id} output is unexpectedly small.")
            text = docx_text(output_path)
            assert_true("Taylor Applicant" in text, f"{template_id} missing sample candidate name.")
            assert_social_contact_line(text, template_id)
            for forbidden in forbidden_reference_terms:
                assert_true(forbidden not in text, f"{template_id} leaked reference template content: {forbidden}")

        no_photo_path = tmp_path / "visual_photo_optional_no_photo.docx"
        no_photo_result = render_cv_with_docx_template(
            structured_cv=cv,
            template_id="visual_photo_optional",
            output_path=str(no_photo_path),
            metadata={"smoke_test": True},
        )
        assert_true(no_photo_result["success"] is True, f"Photo template without photo failed: {no_photo_result}")
        assert_true(docx_has_media(no_photo_path) is False, "Photo template without photo should not embed media.")

        with_photo_path = tmp_path / "visual_photo_optional_with_photo.docx"
        with_photo_result = render_cv_with_docx_template(
            structured_cv=cv,
            template_id="visual_photo_optional",
            output_path=str(with_photo_path),
            metadata={"smoke_test": True},
            photo_bytes=sample_photo_bytes(),
            photo_filename="sample_photo.png",
        )
        assert_true(with_photo_result["success"] is True, f"Photo template with photo failed: {with_photo_result}")
        assert_true(docx_has_media(with_photo_path) is True, "Photo template with photo should embed media.")
        assert_true(first_docx_media_size(with_photo_path) == (512, 512), "Square photo should be embedded as prepared square PNG.")

        portrait_photo_path = tmp_path / "visual_photo_optional_portrait_photo.docx"
        portrait_result = render_cv_with_docx_template(
            structured_cv=cv,
            template_id="visual_photo_optional",
            output_path=str(portrait_photo_path),
            metadata={"smoke_test": True},
            photo_bytes=sample_photo_bytes(120, 220),
            photo_filename="portrait_sample.png",
        )
        assert_true(portrait_result["success"] is True, f"Portrait photo render failed: {portrait_result}")
        assert_true(first_docx_media_size(portrait_photo_path) == (512, 512), "Portrait photo should be center-cropped without distortion.")

        landscape_photo_path = tmp_path / "visual_photo_optional_landscape_photo.docx"
        landscape_result = render_cv_with_docx_template(
            structured_cv=cv,
            template_id="visual_photo_optional",
            output_path=str(landscape_photo_path),
            metadata={"smoke_test": True},
            photo_bytes=sample_photo_bytes(220, 120),
            photo_filename="landscape_sample.png",
        )
        assert_true(landscape_result["success"] is True, f"Landscape photo render failed: {landscape_result}")
        assert_true(first_docx_media_size(landscape_photo_path) == (512, 512), "Landscape photo should be center-cropped without distortion.")

        client = TestClient(app)
        route_response = client.post(
            "/ats-cv/export-docx",
            data={
                "ats_cv_json": json.dumps(cv, ensure_ascii=False),
                "template_id": "visual_photo_optional",
                "language": "English",
                "one_page": "false",
                "enabled_sections": "",
                "export_style": "standard",
                "docx_render_mode": "template",
                "docx_template_id": "visual_photo_optional",
                "include_photo": "true",
            },
            files={"cv_photo": ("sample_photo.png", sample_photo_bytes(120, 220), "image/png")},
        )
        assert_true(route_response.status_code == 200, f"Photo export endpoint failed: {route_response.status_code} {route_response.text}")
        route_docx_path = tmp_path / "route_photo_export.docx"
        route_docx_path.write_bytes(route_response.content)
        assert_true(docx_has_media(route_docx_path) is True, "Photo export endpoint should embed media.")

        programmatic_docx_response = client.post(
            "/ats-cv/export-docx",
            data={
                "ats_cv_json": json.dumps(cv, ensure_ascii=False),
                "template_id": "visual_photo_optional",
                "language": "English",
                "one_page": "false",
                "enabled_sections": "",
                "export_style": "standard",
                "docx_render_mode": "programmatic",
                "include_photo": "true",
            },
            files={"cv_photo": ("sample_photo.png", sample_photo_bytes(220, 120), "image/png")},
        )
        assert_true(
            programmatic_docx_response.status_code == 200,
            f"Programmatic photo DOCX endpoint failed: {programmatic_docx_response.status_code} {programmatic_docx_response.text}",
        )
        programmatic_docx_path = tmp_path / "programmatic_photo_export.docx"
        programmatic_docx_path.write_bytes(programmatic_docx_response.content)
        assert_true(docx_has_media(programmatic_docx_path) is True, "Programmatic DOCX photo export should embed media.")

        pdf_response = client.post(
            "/ats-cv/export-pdf",
            data={
                "ats_cv_json": json.dumps(cv, ensure_ascii=False),
                "template_id": "visual_photo_optional",
                "language": "English",
                "one_page": "false",
                "enabled_sections": "",
                "export_style": "standard",
                "include_photo": "true",
            },
            files={"cv_photo": ("sample_photo.png", sample_photo_bytes(220, 120), "image/png")},
        )
        assert_true(pdf_response.status_code == 200, f"Photo PDF endpoint failed: {pdf_response.status_code} {pdf_response.text}")
        assert_true(pdf_response.content.startswith(b"%PDF"), "Photo PDF endpoint should return a PDF.")
        assert_true(len(pdf_response.content) > 3_000, "Photo PDF endpoint returned an unexpectedly small file.")

        prompt = build_ats_cv_generation_prompt(
            cv_text="BS Computer Engineering, Example University, 2021-2025.",
            job_description="Backend developer role.",
            template=next(template for template in get_ats_cv_templates() if template["id"] == "modern_professional"),
            language="English",
            adaptation_level="balanced",
        )
        assert_true("Do not infer or add \"student\"" in prompt, "ATS CV prompt should prevent unsupported student wording.")

        guarded = _neutralize_unsupported_student_wording(
            {"professional_summary": "Computer Engineering student with API experience."},
            "BS Computer Engineering, Example University, 2021-2025.",
            "Backend developer role.",
        )
        assert_true("student" not in guarded["professional_summary"].lower(), "Student wording guard should neutralize unsupported student claims.")

    print("visual cv template smoke: ok")


if __name__ == "__main__":
    main()
