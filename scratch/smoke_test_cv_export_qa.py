import json
import os
import sys
import uuid
import re
from io import BytesIO
from datetime import datetime
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from fastapi.testclient import TestClient
from main import app
from models import JobApplicationAsset
from services.job_application_asset_service import serialize_asset
from services.docx_template_service import render_cv_with_docx_template, get_docx_template_catalog
from services.ats_cv_schema import get_empty_ats_cv_schema
from services.ats_cv_export_service import build_plain_text_preview, balance_one_page_content
from services.ats_cv_templates import get_ats_cv_template
from services.ats_cv_schema import normalize_structured_cv_schema_fields
from services.ats_cv_adaptation import standardize_cv_adaptation_quality
from docx import Document
from pypdf import PdfReader

client = TestClient(app)


def assert_true(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def valid_mock_cv() -> dict:
    cv = get_empty_ats_cv_schema()
    cv["contact"] = {
        "full_name": "Alex Candidate",
        "target_title": "Junior Backend Developer",
        "email": "alex@example.com",
        "phone": "+1 555 555 5555",
        "location": "Remote",
        "linkedin": "https://linkedin.com/in/alexcandidate",
        "github": "https://github.com/alexcandidate",
        "portfolio": ""
    }
    cv["professional_summary"] = "Junior backend developer with project-based Python, API, SQL, testing, and documentation experience."
    cv["career_objective"] = "To build reliable backend services and APIs."
    cv["technical_summary"] = "Python, SQL, REST APIs, FastAPI, Git, Docker."
    cv["skills"] = {
        "technical_skills": ["Python", "SQL"],
        "core_skills": ["REST APIs", "FastAPI"],
        "tools": ["Git", "Docker"],
        "databases": [],
        "cloud": [],
        "soft_skills": []
    }
    cv["experience"] = [
        {
            "company": "Example Software Studio",
            "role": "Backend Developer Intern",
            "location": "Remote",
            "start_date": "2025",
            "end_date": "2026",
            "bullets": [
                "Built REST API endpoints for internal demo services.",
                "Documented testing notes and integration behavior.",
            ],
        }
    ]
    cv["projects"] = [
        {
            "name": "API Tracker",
            "description": "Built a small API tracking project.",
            "technologies": ["Python", "FastAPI", "SQLite"],
            "bullets": ["Implemented endpoints, validation logic, and SQL queries."],
            "link": ""
        }
    ]
    cv["education"] = [
        {
            "school": "Example University",
            "degree": "BS Computer Engineering",
            "department": "",
            "start_date": "2021",
            "end_date": "2025",
            "details": []
        }
    ]
    cv["certifications"] = [{"name": "Python Foundations", "issuer": "Example Academy", "date": "2025", "link": ""}]
    cv["languages"] = [{"language": "English", "level": "Professional"}]
    cv["ats_metadata"] = {
        "target_role": "Junior Backend Developer",
        "target_company": "",
        "job_family": "Technology",
        "job_keywords_used": ["Python", "SQL"],
        "transferable_keywords_used": [],
        "missing_keywords": [],
        "risky_keywords_not_added": [],
        "ats_score_before": 50,
        "ats_score_after": 85,
        "optimization_summary": "Added tech stack detail.",
        "alignment_confidence": "high",
        "adaptation_notes": [],
        "ats_score_explanation": {
            "before_reason": "Minimal content.",
            "after_reason": "Included structured role keywords.",
            "improvement_reasons": [],
            "remaining_gaps": []
        }
    }
    return cv


def rich_mock_cv() -> dict:
    cv = valid_mock_cv()
    cv["experience"] = [
        {
            "company": "Example Bank",
            "role": "Payment Systems Backend Intern",
            "location": "Istanbul",
            "start_date": "2025",
            "end_date": "2026",
            "bullets": [
                "Built and tested backend service endpoints for payment operations workflows.",
                "Documented API behavior, release notes, and integration test scenarios.",
                "Reviewed payment transaction logs to support troubleshooting and validation.",
                "Collaborated with senior developers on SQL checks and deployment readiness.",
                "Mapped business requirements into implementation notes for backend tasks.",
            ],
        },
        {
            "company": "Example Retail",
            "role": "Customer Operations Associate",
            "location": "Remote",
            "start_date": "2024",
            "end_date": "2025",
            "bullets": [
                "Resolved customer requests in a time-sensitive service environment.",
                "Documented recurring issues and escalated process gaps.",
                "Coordinated with teammates to maintain accurate operational records.",
                "Supported onboarding notes for new team members.",
            ],
        },
    ]
    cv["projects"] = [
        {
            "name": "Payment API Monitor",
            "description": "Built a backend monitoring tool for API checks and operational reporting.",
            "technologies": ["Python", "FastAPI", "PostgreSQL", "Docker"],
            "bullets": [
                "Implemented endpoints for transaction status review.",
                "Added SQL queries for validation and issue triage.",
                "Created concise test notes and technical documentation.",
            ],
            "link": "",
        },
        {
            "name": "Application Materials Portal",
            "description": "Built a document generation prototype.",
            "technologies": ["Python", "FastAPI", "SQLite"],
            "bullets": [
                "Implemented validation and preview behavior.",
                "Added export checks for generated application documents.",
            ],
            "link": "",
        },
    ]
    cv["skills"]["technical_skills"] = ["Python", "FastAPI", "REST APIs", "SQL", "PostgreSQL", "Testing"]
    cv["skills"]["soft_skills"] = ["Documentation", "Communication", "Teamwork", "Customer focus", "Time management"]
    cv["certifications"] = [
        {"name": "Python Foundations", "issuer": "Example Academy", "date": "2025", "link": ""},
        {"name": "SQL Basics", "issuer": "Example Academy", "date": "2025", "link": ""},
        {"name": "Cloud Intro", "issuer": "Example Academy", "date": "2024", "link": ""},
    ]
    return cv


def project_integrity_cv() -> dict:
    cv = valid_mock_cv()
    cv["contact"]["target_title"] = "Backend Developer"
    cv["professional_summary"] = "Backend developer with payment, merchant integration, and full-stack project experience."
    cv["career_objective"] = "Build reliable backend services for payment and merchant workflows."
    cv["ats_metadata"]["target_role"] = "Merchant Payment Backend Developer"
    cv["ats_metadata"]["job_family"] = "fintech_payment"
    cv["ats_metadata"]["job_keywords_used"] = [
        "merchant",
        "external API",
        "file upload",
        "validation",
        "payment",
    ]
    cv["experience"] = [
        {
            "company": "Inditex Zara",
            "role": "Sales Associate",
            "location": "Istanbul",
            "start_date": "2024",
            "end_date": "2025",
            "bullets": [
                "Handled retail customer requests and store operations.",
                "Supported stock organization and in-store communication.",
                "Documented recurring customer service issues.",
            ],
        },
        {
            "company": "VakifBank",
            "role": "Payment Systems Intern",
            "location": "Istanbul",
            "start_date": "2025",
            "end_date": "2026",
            "bullets": [
                "Supported payment transaction validation and backend workflow documentation.",
                "Reviewed merchant payment scenarios with API and SQL checks.",
                "Documented banking integration test notes for payment services.",
            ],
        },
    ]
    cv["projects"] = [
        {
            "name": "Restaurant POS System - Full-Stack Project",
            "description": "Restaurant table order payment management for dining operations.",
            "technologies": ["ASP.NET Core", "Supabase PostgreSQL", "JWT"],
            "bullets": [
                "Implemented restaurant table order and payment workflows.",
                "Built POS screens for menu, table, and receipt management.",
                "Added JWT authentication for restaurant staff access.",
            ],
            "link": "",
        },
        {
            "name": "Merchant Application System - Web API Project",
            "description": "Merchant external service integration with validation and upload workflows.",
            "technologies": ["Entity Framework Core", "FluentValidation", "Google Maps API"],
            "bullets": [
                "Implemented merchant external API integration endpoints.",
                "Added file upload validation for merchant application documents.",
                "Used FluentValidation rules for merchant onboarding requests.",
            ],
            "link": "",
        },
    ]
    return cv


def fraud_payment_source_cv() -> dict:
    cv = valid_mock_cv()
    cv["contact"]["target_title"] = "Junior Backend Developer"
    cv["professional_summary"] = "Junior backend developer with API, SQL, documentation, and project-based implementation experience."
    cv["career_objective"] = "Build reliable backend services and APIs."
    cv["skills"] = {
        "technical_skills": ["C#", "ASP.NET Core", "Python", "SQL", "REST APIs"],
        "core_skills": ["Backend Development", "Requirement Analysis"],
        "tools": ["Git", "Postman", "MS Office"],
        "databases": ["PostgreSQL", "SQL Server"],
        "cloud": [],
        "soft_skills": ["Communication", "Attention to Detail", "Problem Solving"],
    }
    cv["experience"] = [
        {
            "company": "VakifBank",
            "role": "Payment Systems Application Development Intern",
            "location": "Istanbul",
            "start_date": "2025",
            "end_date": "2026",
            "bullets": [
                "Supported payment transaction validation and backend workflow documentation.",
                "Reviewed API test scenarios with Postman and SQL checks.",
                "Documented business requirements for payment systems processes.",
                "Implemented backend endpoint updates for internal application workflows.",
            ],
        },
        {
            "company": "Inditex Zara",
            "role": "Sales Associate",
            "location": "Istanbul",
            "start_date": "2024",
            "end_date": "2025",
            "bullets": [
                "Handled high-volume customer interactions with attention to detail.",
                "Supported operational consistency across store tasks.",
                "Communicated recurring issues to teammates and supervisors.",
            ],
        },
    ]
    cv["projects"] = [
        {
            "name": "Merchant Application System - Web API Project",
            "description": "Merchant external service integration with validation and upload workflows.",
            "technologies": ["ASP.NET Core", "Entity Framework Core", "FluentValidation", "SQL Server"],
            "bullets": [
                "Implemented merchant external API integration endpoints.",
                "Added file upload validation for merchant application documents.",
                "Used FluentValidation rules for data integrity checks.",
            ],
            "link": "",
        },
        {
            "name": "Restaurant POS System",
            "description": "Restaurant table order and payment management project.",
            "technologies": ["ASP.NET Core", "PostgreSQL", "JWT"],
            "bullets": [
                "Built order and payment workflow screens.",
                "Added role-based authentication with JWT.",
            ],
            "link": "",
        },
    ]
    cv["ats_metadata"] = {
        "target_role": "Fraud Operations Analyst",
        "target_company": "",
        "job_family": "",
        "job_keywords_used": [],
        "transferable_keywords_used": [],
        "missing_keywords": [],
        "risky_keywords_not_added": [],
        "ats_score_before": 50,
        "ats_score_after": 78,
        "optimization_summary": "",
        "alignment_confidence": "medium",
        "adaptation_notes": [],
        "ats_score_explanation": {
            "before_reason": "",
            "after_reason": "",
            "improvement_reasons": [],
            "remaining_gaps": [],
        },
    }
    return cv


def fraud_risk_payment_job_description() -> str:
    return (
        "Junior Fraud and Risk Operations Analyst role for an e-commerce payments team. "
        "Responsibilities include reviewing suspicious payment transactions, supporting fraud prevention operations, "
        "documenting risk workflows, validating transaction data, using SQL and Excel for data analysis, "
        "checking API/payment integration issues, and communicating findings with operations and product teams."
    )


# Recreate Streamlit's local helper functions to verify metadata extraction works without crash
def unwrap_asset_structured_json(asset: dict) -> dict:
    structured = asset.get("structured_json") if isinstance(asset, dict) else {}
    return structured if isinstance(structured, dict) else {}


def get_asset_quality_report(asset: dict) -> dict:
    structured = unwrap_asset_structured_json(asset)
    return structured.get("quality_report") if isinstance(structured.get("quality_report"), dict) else {}


def get_asset_structure_report(asset: dict) -> dict:
    structured = unwrap_asset_structured_json(asset)
    return structured.get("structure_report") if isinstance(structured.get("structure_report"), dict) else {}


def test_ats_cv_builder_exports() -> None:
    cv = valid_mock_cv()
    payload = {
        "ats_cv_json": json.dumps(cv, ensure_ascii=False),
        "template_id": "classic_ats",
        "language": "English",
        "one_page": "false",
        "enabled_sections": "",
        "export_style": "standard",
    }

    # 1. Programmatic DOCX
    p_docx_payload = dict(payload)
    p_docx_payload["docx_render_mode"] = "programmatic"
    res = client.post("/ats-cv/export-docx", data=p_docx_payload)
    assert_true(res.status_code == 200, f"Programmatic DOCX failed: {res.status_code}")
    assert_true(len(res.content) > 0, "Programmatic DOCX empty")

    # 2. Template DOCX - Classic
    t_classic_payload = dict(payload)
    t_classic_payload["docx_render_mode"] = "template"
    t_classic_payload["docx_template_id"] = "ats_classic_docx"
    res = client.post("/ats-cv/export-docx", data=t_classic_payload)
    assert_true(res.status_code == 200, f"Template DOCX Classic failed: {res.status_code}")
    assert_true(len(res.content) > 0, "Template DOCX Classic empty")

    # 3. Template DOCX - Modern
    t_modern_payload = dict(payload)
    t_modern_payload["docx_render_mode"] = "template"
    t_modern_payload["docx_template_id"] = "ats_modern_docx"
    res = client.post("/ats-cv/export-docx", data=t_modern_payload)
    assert_true(res.status_code == 200, f"Template DOCX Modern failed: {res.status_code}")
    assert_true(len(res.content) > 0, "Template DOCX Modern empty")

    # 4. PDF Export (offline supported via ReportLab)
    res = client.post("/ats-cv/export-pdf", data=payload)
    assert_true(res.status_code == 200, f"PDF export failed: {res.status_code}")
    assert_true(len(res.content) > 0, "PDF export empty")

    # 5. TXT Export
    res = client.post("/ats-cv/export-txt", data=payload)
    assert_true(res.status_code == 200, f"TXT export failed: {res.status_code}")
    assert_true(len(res.content) > 0, "TXT export empty")
    txt = res.content.decode("utf-8")
    assert_true("SUMMARY" in txt, "English TXT export should use SUMMARY heading.")
    assert_true("PROFESSIONAL SUMMARY" not in txt, "English TXT export should not use PROFESSIONAL SUMMARY heading.")

    tr_payload = dict(payload)
    tr_payload["language"] = "Turkish"
    res = client.post("/ats-cv/export-txt", data=tr_payload)
    assert_true(res.status_code == 200, f"Turkish TXT export failed: {res.status_code}")
    tr_txt = res.content.decode("utf-8")
    assert_true("ÖZET" in tr_txt, "Turkish TXT export should use ÖZET heading.")

    # 6. Unknown Template ID returns clean error
    bad_payload = dict(payload)
    bad_payload["docx_render_mode"] = "template"
    bad_payload["docx_template_id"] = "unknown_template_style"
    res = client.post("/ats-cv/export-docx", data=bad_payload)
    assert_true(res.status_code == 400, f"Expected 400 on unknown template, got {res.status_code}")
    assert_true("fallback" in res.json()["detail"].lower(), "Traceback/unclean error exposed")


def test_one_page_preserves_priority_content() -> None:
    cv = rich_mock_cv()
    template = get_ats_cv_template("modern_professional")
    normal_text = build_plain_text_preview(cv, template, "English", export_style="standard")
    one_page_cv = balance_one_page_content(cv, template, "English")
    one_page_text = build_plain_text_preview(one_page_cv, template, "English", export_style="balanced_one_page")

    assert_true(normal_text.count("- ") > one_page_text.count("- "), "Normal export should keep fuller bullet content.")
    assert_true(len(one_page_cv["experience"][0]["bullets"]) >= 3, "Primary experience should keep at least 3 bullets.")
    assert_true(len(one_page_cv["experience"][1]["bullets"]) >= 2, "Secondary experience should keep at least 2 bullets.")
    assert_true(len(one_page_cv["projects"][0]["bullets"]) >= 2, "Main project should keep at least 2 bullets.")
    assert_true(len(one_page_cv["projects"]) >= 2, "One-page mode should preserve at least two projects when available.")
    assert_true("SUMMARY" in one_page_text, "One-page preview should use SUMMARY heading.")
    assert_true("PROFESSIONAL SUMMARY" not in one_page_text, "One-page preview should not use PROFESSIONAL SUMMARY heading.")


def test_schema_summary_alias_exports() -> None:
    base_cv = valid_mock_cv()
    cases = [
        ("summary", {"summary": "Alias summary for export.", "professional_summary": "", "career_objective": ""}),
        ("professional_summary", {"professional_summary": "Professional summary fallback.", "career_objective": ""}),
        ("career_objective", {"career_objective": "Preserved career objective.", "professional_summary": ""}),
    ]

    for case_name, overrides in cases:
        cv = dict(base_cv)
        cv.update(overrides)
        if case_name == "summary":
            cv.pop("career_objective", None)
        if case_name == "professional_summary":
            cv.pop("career_objective", None)

        normalized = normalize_structured_cv_schema_fields(cv)
        assert_true("career_objective" in normalized, f"{case_name} should normalize career_objective.")
        if case_name == "career_objective":
            assert_true(normalized["career_objective"] == "Preserved career objective.", "Existing career_objective should be preserved.")

        payload = {
            "ats_cv_json": json.dumps(cv, ensure_ascii=False),
            "template_id": "modern_professional",
            "language": "English",
            "one_page": "false",
            "enabled_sections": "",
            "export_style": "standard",
            "docx_render_mode": "programmatic",
        }
        for endpoint in ["export-docx", "export-pdf", "export-txt"]:
            res = client.post(f"/ats-cv/{endpoint}", data=payload)
            assert_true(res.status_code == 200, f"{endpoint} failed for {case_name}: {res.status_code} {res.text}")

        txt_res = client.post("/ats-cv/export-txt", data=payload)
        txt = txt_res.content.decode("utf-8")
        assert_true("SUMMARY" in txt, f"{case_name} export should display SUMMARY heading.")
        assert_true("PROFESSIONAL SUMMARY" not in txt, f"{case_name} export should not display PROFESSIONAL SUMMARY.")


def test_one_page_project_and_experience_integrity() -> None:
    cv = project_integrity_cv()

    for template_id in ["visual_photo_optional", "modern_professional"]:
        payload = {
            "ats_cv_json": json.dumps(cv, ensure_ascii=False),
            "template_id": template_id,
            "language": "English",
            "one_page": "true",
            "enabled_sections": "",
            "export_style": "balanced_one_page",
            "docx_render_mode": "programmatic",
        }

        txt_res = client.post("/ats-cv/export-txt", data=payload)
        assert_true(txt_res.status_code == 200, f"TXT integrity export failed for {template_id}: {txt_res.status_code} {txt_res.text}")
        assert_project_and_experience_integrity(txt_res.content.decode("utf-8"), f"{template_id} TXT")

        pdf_res = client.post("/ats-cv/export-pdf", data=payload)
        assert_true(pdf_res.status_code == 200, f"PDF integrity export failed for {template_id}: {pdf_res.status_code} {pdf_res.text}")
        pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_res.content)).pages)
        assert_project_and_experience_integrity(pdf_text, f"{template_id} PDF")

    docx_payload = {
        "ats_cv_json": json.dumps(cv, ensure_ascii=False),
        "template_id": "visual_photo_optional",
        "language": "English",
        "one_page": "true",
        "enabled_sections": "",
        "export_style": "balanced_one_page",
        "docx_render_mode": "programmatic",
    }
    docx_res = client.post("/ats-cv/export-docx", data=docx_payload)
    assert_true(docx_res.status_code == 200, f"DOCX integrity export failed: {docx_res.status_code} {docx_res.text}")
    assert_project_and_experience_integrity(_docx_content_text(docx_res.content), "visual_photo_optional DOCX")


def test_deterministic_adaptation_quality_standard() -> None:
    source_cv = fraud_payment_source_cv()
    job_description = fraud_risk_payment_job_description()

    conservative = standardize_cv_adaptation_quality(
        source_cv,
        job_description=job_description,
        adaptation_level="conservative",
        source_cv_text=_structured_text(source_cv),
        language="English",
    )
    balanced = standardize_cv_adaptation_quality(
        source_cv,
        job_description=job_description,
        adaptation_level="balanced",
        source_cv_text=_structured_text(source_cv),
        language="English",
    )
    strong = standardize_cv_adaptation_quality(
        source_cv,
        job_description=job_description,
        adaptation_level="strong",
        source_cv_text=_structured_text(source_cv),
        language="English",
    )

    strong_text = _structured_text(strong).lower()
    strong_title = strong["contact"]["target_title"].lower()
    strong_summary = strong["professional_summary"].lower()
    strong_skills = _flatten_skill_values(strong["skills"])
    strong_skill_text = " ".join(strong_skills).lower()

    assert_true("fraud" in strong_title or "risk" in strong_title, "Strong title should align to fraud/risk/payment job.")
    assert_true("payment" in strong_title, "Strong title should preserve payment-domain alignment.")
    for term in ["fraud", "risk", "payment", "validation"]:
        assert_true(term in strong_summary, f"Strong summary should include {term} relevance.")
    for term in ["data validation", "api testing", "sql querying", "payment systems"]:
        assert_true(term in strong_skill_text, f"Strong skills should include supported target skill: {term}")

    first_skills = [skill.lower() for skill in strong_skills[:6]]
    assert_true(any("data validation" in skill for skill in first_skills), "Data Validation should be prioritized early.")
    assert_true(any("api testing" in skill for skill in first_skills), "API Testing should be prioritized early.")
    assert_true(any("sql" in skill for skill in first_skills), "SQL Querying should be prioritized early.")

    first_experience = strong["experience"][0]
    assert_true(first_experience["company"] == "VakifBank", "Relevant payment experience should be prioritized.")
    first_bullet_text = " ".join(first_experience["bullets"][:2]).lower()
    assert_true("payment" in first_bullet_text, "Payment evidence should remain high in experience bullets.")
    assert_true(
        any(term in first_bullet_text for term in ["validation", "api", "sql", "documentation"]),
        "Strong experience bullets should prioritize validation/API/SQL/documentation evidence.",
    )

    first_project = strong["projects"][0]
    project_text = _structured_text(first_project).lower()
    assert_true("merchant" in project_text and "validation" in project_text, "Relevant merchant/validation project should be prioritized.")

    unsupported_phrases = [
        "direct fraud investigation",
        "senior fraud",
        "fraud investigation lead",
        "chargeback specialist",
        "aml",
        "kyc",
        "production fraud monitoring",
        "risk rule creation",
    ]
    for phrase in unsupported_phrases:
        assert_true(phrase not in strong_text, f"Strong adaptation invented unsupported claim: {phrase}")

    assert_true(conservative["contact"]["target_title"] == "Junior Backend Developer", "Conservative should preserve source title.")
    assert_true(
        "risk/fraud operations awareness" not in " ".join(_flatten_skill_values(balanced["skills"])).lower(),
        "Balanced should be less aggressive than strong for transferable fraud/risk wording.",
    )
    assert_true(
        "risk/fraud operations awareness" in strong_skill_text,
        "Strong should include cautious transferable fraud/risk awareness when payment/data validation evidence exists.",
    )

    report = strong["ats_metadata"].get("adaptation_quality_report", {})
    assert_true(report.get("adaptation_level") == "strong", "Adaptation report should record strong level.")
    assert_true(report.get("detected_domain") == "fraud_risk_payments", "Adaptation report should detect fraud/risk/payments domain.")
    assert_true(isinstance(report.get("warnings"), list), "Adaptation report warnings should be a list.")

    for proper_noun in ["VakifBank", "Inditex Zara", "Merchant Application System - Web API Project"]:
        assert_true(proper_noun in _structured_text(strong), f"Proper noun was corrupted or removed: {proper_noun}")

    template = get_ats_cv_template("modern_professional")
    rendered_text = build_plain_text_preview(strong, template, "English", export_style="standard")
    assert_true("SUMMARY" in rendered_text, "Template rendering should preserve SUMMARY heading.")
    assert_true("PROFESSIONAL SUMMARY" not in rendered_text, "Template rendering should not revert to PROFESSIONAL SUMMARY.")
    assert_true("fraud" in rendered_text.lower() and "payment" in rendered_text.lower(), "Template rendering should not weaken strong adaptation.")

    one_page_cv = balance_one_page_content(strong, template, "English")
    one_page_text = _structured_text(one_page_cv).lower()
    assert_true("payment" in one_page_text, "One-page mode should preserve payment relevance.")
    assert_true(
        any(term in one_page_text for term in ["validation", "api", "sql"]),
        "One-page mode should preserve validation/API/SQL relevance.",
    )
    assert_true(one_page_cv["experience"][0]["company"] == "VakifBank", "One-page mode should preserve relevant experience priority.")
    assert_true("Merchant Application System - Web API Project" in _structured_text(one_page_cv), "One-page mode should preserve relevant project identity.")


def assert_project_and_experience_integrity(text: str, context: str) -> None:
    normalized = " ".join(str(text or "").split())

    restaurant = _segment_for_title(
        normalized,
        "Restaurant POS System - Full-Stack Project",
        ["Merchant Application System - Web API Project", "EDUCATION", "CERTIFICATIONS", "LANGUAGES"],
    )
    merchant = _segment_for_title(
        normalized,
        "Merchant Application System - Web API Project",
        ["Restaurant POS System - Full-Stack Project", "EDUCATION", "CERTIFICATIONS", "LANGUAGES"],
    )
    assert_true("Restaurant table order payment management" in restaurant, f"{context}: Restaurant description moved or missing.")
    assert_true("Supabase PostgreSQL" in restaurant and "JWT" in restaurant, f"{context}: Restaurant technologies moved or missing.")
    assert_true("restaurant table order" in restaurant.lower(), f"{context}: Restaurant bullet moved or missing.")
    assert_true("Merchant external service integration" not in restaurant, f"{context}: Merchant description appears under Restaurant.")
    assert_true("Entity Framework Core" not in restaurant and "FluentValidation" not in restaurant, f"{context}: Merchant technologies appear under Restaurant.")

    assert_true("Merchant external service integration" in merchant, f"{context}: Merchant description moved or missing.")
    assert_true("Entity Framework Core" in merchant and "FluentValidation" in merchant, f"{context}: Merchant technologies moved or missing.")
    assert_true("merchant external api integration" in merchant.lower(), f"{context}: Merchant bullet moved or missing.")
    assert_true("Restaurant table order payment management" not in merchant, f"{context}: Restaurant description appears under Merchant.")
    assert_true("Supabase PostgreSQL" not in merchant and "JWT" not in merchant, f"{context}: Restaurant technologies appear under Merchant.")

    vakifbank = _segment_for_title(
        normalized,
        "VakifBank",
        ["Inditex Zara", "PROJECTS", "EDUCATION", "CERTIFICATIONS", "LANGUAGES"],
    )
    zara = _segment_for_title(
        normalized,
        "Inditex Zara",
        ["VakifBank", "PROJECTS", "EDUCATION", "CERTIFICATIONS", "LANGUAGES"],
    )
    assert_true("payment transaction validation" in vakifbank.lower(), f"{context}: VakifBank bullet moved or missing.")
    assert_true("retail customer requests" not in vakifbank.lower(), f"{context}: Zara bullet appears under VakifBank.")
    assert_true("retail customer requests" in zara.lower(), f"{context}: Zara bullet moved or missing.")
    assert_true("payment transaction validation" not in zara.lower(), f"{context}: VakifBank bullet appears under Zara.")


def _segment_for_title(text: str, title: str, next_markers: list[str]) -> str:
    start = text.find(title)
    assert_true(start >= 0, f"Missing title in rendered output: {title}")
    end_candidates = [text.find(marker, start + len(title)) for marker in next_markers]
    end_candidates = [candidate for candidate in end_candidates if candidate >= 0]
    end = min(end_candidates) if end_candidates else len(text)
    return text[start:end]


def _docx_content_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    lines = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                lines.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(lines)


def _flatten_skill_values(skills: dict) -> list[str]:
    values = []
    if isinstance(skills, dict):
        for group_values in skills.values():
            if isinstance(group_values, list):
                values.extend(str(value) for value in group_values if str(value or "").strip())
    return values


def _structured_text(value) -> str:
    if isinstance(value, dict):
        return " ".join(_structured_text(item) for item in value.values())
    if isinstance(value, list):
        return " ".join(_structured_text(item) for item in value)
    return str(value or "")


def test_generated_asset_compatibility() -> None:
    # Verify we can serialize an asset and that helpers extract report safely.
    # 1. Simulate an asset with full metadata
    meta = {
        "quality_report": {
            "quality_score": 90,
            "issues": [],
            "critical_count": 0,
            "warning_count": 0,
            "info_count": 0,
            "summary": "Looks clean."
        },
        "structure_report": {
            "structure_score": 95,
            "issues": []
        },
        "adaptation_level": "strong",
        "docx_render_mode": "template",
        "docx_template_id": "ats_modern_docx"
    }

    mock_db_asset = JobApplicationAsset(
        id=123,
        job_id=1,
        asset_type="tailored_cv",
        title="Tailored CV",
        content_text="Mock plain text content",
        structured_json=json.dumps(meta),
        file_path="generated_assets/tailored_cv_ats_modern_docx_20260624_120000.docx",
        export_format="docx",
        template_id="ats_modern_docx",
        language="English",
        created_at=datetime.utcnow(),
        updated_at=datetime.utcnow()
    )

    serialized = serialize_asset(mock_db_asset)
    assert_true(serialized["structured_json"] == meta, "Metadata serialization mismatch")

    # Test helpers on new asset dict
    q_rep = get_asset_quality_report(serialized)
    s_rep = get_asset_structure_report(serialized)
    assert_true(q_rep.get("quality_score") == 90, "Failed to get quality score from metadata")
    assert_true(s_rep.get("structure_score") == 95, "Failed to get structure score from metadata")

    # 2. Simulate old asset without new metadata keys
    old_meta = {
        "target_role": "Backend Engineer",
        "ats_score_before": 45
    }
    mock_db_asset_old = JobApplicationAsset(
        id=124,
        job_id=1,
        asset_type="tailored_cv",
        title="Old Tailored CV",
        content_text="Mock old content",
        structured_json=json.dumps(old_meta),
        file_path="generated_assets/tailored_cv_old.docx",
        export_format="docx",
        created_at=datetime.utcnow()
    )
    serialized_old = serialize_asset(mock_db_asset_old)
    q_rep_old = get_asset_quality_report(serialized_old)
    s_rep_old = get_asset_structure_report(serialized_old)
    assert_true(q_rep_old == {}, "Old asset quality report should be empty dictionary")
    assert_true(s_rep_old == {}, "Old asset structure report should be empty dictionary")

    # 3. Simulate asset with invalid JSON string in db
    mock_db_asset_invalid = JobApplicationAsset(
        id=125,
        job_id=1,
        asset_type="tailored_cv",
        structured_json="invalid { json",
        created_at=datetime.utcnow()
    )
    serialized_invalid = serialize_asset(mock_db_asset_invalid)
    assert_true(serialized_invalid["structured_json"] is None, "Invalid JSON should serialize to None")
    assert_true(get_asset_quality_report(serialized_invalid) == {}, "Invalid JSON quality report should be empty")
    assert_true(get_asset_structure_report(serialized_invalid) == {}, "Invalid JSON structure report should be empty")


def test_filename_validation() -> None:
    # Import filename generator helpers
    from services.job_application_asset_service import _cv_asset_filename
    from routers.ats_cv_builder import _cv_export_filename

    names = [
        _cv_asset_filename("Tailored CV", "Ats Modern Docx", "docx"),
        _cv_export_filename("ATS CV", "Classic ATS", "pdf")
    ]

    for name in names:
        assert_true(name == name.lower(), f"Filename must be lowercase: {name}")
        assert_true(" " not in name, f"Filename must not contain spaces: {name}")
        # check pattern
        assert_true(name.startswith("tailored_cv_") or name.startswith("ats_cv_"), f"Unexpected prefix: {name}")
        # check safe chars
        assert_true(re.match(r"^[a-z0-9_\.]+$", name) is not None, f"Invalid characters in filename: {name}")
        # check no personal name
        assert_true("alex" not in name and "candidate" not in name, f"Filename contains candidate name: {name}")


def test_file_content_sanity() -> None:
    # Render with docx templates to temporary files, and check their content
    os.makedirs("scratch", exist_ok=True)
    cv = valid_mock_cv()
    
    classic_temp = f"scratch/test_sanity_classic_{uuid.uuid4().hex}.docx"
    modern_temp = f"scratch/test_sanity_modern_{uuid.uuid4().hex}.docx"

    try:
        # Render Classic
        res_classic = render_cv_with_docx_template(
            structured_cv=cv,
            template_id="ats_classic_docx",
            output_path=classic_temp
        )
        assert_true(res_classic["success"], f"Classic render failed: {res_classic}")

        # Render Modern
        res_modern = render_cv_with_docx_template(
            structured_cv=cv,
            template_id="ats_modern_docx",
            output_path=modern_temp
        )
        assert_true(res_modern["success"], f"Modern render failed: {res_modern}")

        # Read contents with python-docx and do sanity checks
        for temp_file in (classic_temp, modern_temp):
            doc = Document(temp_file)
            full_text = []
            for p in doc.paragraphs:
                full_text.append(p.text)
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            combined_text = "\n".join(full_text)
            
            # 1. Non-empty readable text
            assert_true(len(combined_text.strip()) > 50, f"File {temp_file} has insufficient text")
            assert_true("Alex Candidate" in combined_text, f"Candidate name missing from {temp_file}")
            
            # 2. Contact fields are not character-spaced
            assert_true("A l e x" not in combined_text, f"Contact field is character-spaced: {combined_text}")
            
            # 3. LinkedIn/GitHub remain exact
            assert_true("https://linkedin.com/in/alexcandidate" in combined_text, f"LinkedIn URL corrupted: {combined_text}")
            assert_true("https://github.com/alexcandidate" in combined_text, f"GitHub URL corrupted: {combined_text}")
            
            # 4. No empty repeated section headings
            # Section headers should not be repeated or blank
            headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading") or p.text.isupper()]
            non_empty_headings = [h.strip() for h in headings if h.strip()]
            assert_true(len(headings) == len(non_empty_headings), f"Found empty section headings: {headings}")

    finally:
        for f in (classic_temp, modern_temp):
            if os.path.exists(f):
                os.remove(f)


if __name__ == "__main__":
    print("Running QA Export and Preview tests...")
    test_ats_cv_builder_exports()
    print("ATS CV Builder exports: PASS")
    test_one_page_preserves_priority_content()
    print("One-page priority preservation: PASS")
    test_schema_summary_alias_exports()
    print("Schema summary alias exports: PASS")
    test_one_page_project_and_experience_integrity()
    print("One-page project/experience integrity: PASS")
    test_deterministic_adaptation_quality_standard()
    print("Deterministic adaptation quality standard: PASS")
    test_generated_asset_compatibility()
    print("Generated asset compatibility: PASS")
    test_filename_validation()
    print("Filename validation: PASS")
    test_file_content_sanity()
    print("File content sanity checks: PASS")
    print("\nAll QA export and preview stability tests completed successfully!")
